# brd_generator.py - COMPLETE FILE WITH ALL ENHANCEMENTS
# Includes: BRD generation, validation, business rules extraction, API contracts

import os
import json
import re
from typing import Dict, List, Any
from openai import AzureOpenAI
from docx import Document
from io import BytesIO
from dotenv import load_dotenv
from prompts import (
    BRD_SYSTEM_PROMPT,
    COMPLEXITY_ANALYSIS_PROMPT,
    BUSINESS_PROCESS_FLOW_PROMPT,
    POWER_PLATFORM_MAPPING_PROMPT,
    USER_STORY_GENERATION_PROMPT,
    QNA_SYSTEM_PROMPT
)

load_dotenv()

client = AzureOpenAI(
    api_key=os.getenv("AZURE_OPENAI_API_KEY"),
    azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT"),
    api_version="2024-12-01-preview"
)
CHAT_DEPLOY = os.getenv("AZURE_OPENAI_DEPLOYMENT")


def _chat(messages, temperature: float = 0.1):
    """Helper function to call Azure OpenAI"""
    resp = client.chat.completions.create(
        model=CHAT_DEPLOY,
        messages=messages,
        temperature=temperature,
    )
    return resp.choices[0].message.content


def generate_word_brd(content: str, name: str) -> BytesIO:
    """
    Generate a Word document from BRD content (Markdown style)
    """
    doc = Document()
    
    # Add title
    doc.add_heading(f"{name}", 0)
    
    lines = content.split("\n")
    for line in lines:
        line = line.strip()
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("|"):  # Table detection
            # Simple table parsing
            cells = [c.strip() for c in line.split("|") if c.strip()]
            if cells:
                if not hasattr(doc, "_current_table"):
                    table = doc.add_table(rows=1, cols=len(cells))
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    for i, cell in enumerate(cells):
                        hdr_cells[i].text = cell
                    doc._current_table = table
                else:
                    row_cells = doc._current_table.add_row().cells
                    for i, cell in enumerate(cells):
                        row_cells[i].text = cell
        else:
            doc.add_paragraph(line)
            if hasattr(doc, "_current_table"):
                delattr(doc, "_current_table")
    
    # Save to in-memory buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def summarize_graph_enhanced(nodes: List[Dict[str, Any]]) -> str:
    """Enhanced graph summary with detailed component analysis"""
    counts = {}
    complexity_by_type = {}
    
    for n in nodes or []:
        kind = n["kind"]
        counts[kind] = counts.get(kind, 0) + 1
        
        # Aggregate complexity metrics if available
        if "props" in n and "complexity" in n.get("props", {}):
            if kind not in complexity_by_type:
                complexity_by_type[kind] = []
            complexity_by_type[kind].append(n["props"]["complexity"])
    
    lines = [f"- {k}: {v}" for k, v in sorted(counts.items())]
    
    # Add complexity summary
    if complexity_by_type:
        lines.append("\nComplexity Analysis:")
        for kind, complexities in complexity_by_type.items():
            avg_complexity = sum(complexities) / len(complexities)
            lines.append(f"- {kind}: avg complexity {avg_complexity:.1f}")
    
    return "Graph summary (by kind):\n" + "\n".join(lines)


def format_metrics_summary(metrics: Dict[str, Any]) -> str:
    """Format code metrics for inclusion in prompts"""
    total = metrics.get("total", {})
    by_file = metrics.get("by_file", {})
    
    summary = f"""
CODE METRICS SUMMARY:
- Total Files: {total.get('total_files', 0)}
- Total Lines of Code: {total.get('total_loc', 0):,}
- Average Complexity: {total.get('total_complexity', 0) / max(total.get('total_files', 1), 1):.1f}
- Average Maintainability: {total.get('avg_maintainability', 0)}/100

FILE TYPE BREAKDOWN:
"""
    
    for ext, count in total.get('file_types', {}).items():
        summary += f"- {ext}: {count} files\n"
    
    # Add high-risk files
    high_risk_files = []
    for file_path, file_metrics in by_file.items():
        if (file_metrics.get('cyclomatic_complexity', 0) > 15 or 
            file_metrics.get('maintainability_index', 100) < 50):
            high_risk_files.append({
                'file': file_path,
                'complexity': file_metrics.get('cyclomatic_complexity', 0),
                'maintainability': file_metrics.get('maintainability_index', 0),
                'loc': file_metrics.get('lines_of_code', 0)
            })
    
    if high_risk_files:
        summary += "\nHIGH-RISK FILES (complexity > 15 OR maintainability < 50):\n"
        for risk_file in sorted(high_risk_files, key=lambda x: x['complexity'], reverse=True)[:10]:
            summary += f"- {risk_file['file']}: complexity={risk_file['complexity']}, maintainability={risk_file['maintainability']}, LOC={risk_file['loc']}\n"
    
    return summary


def format_business_processes(processes: List[Dict[str, Any]]) -> str:
    """Format business process analysis for prompts"""
    if not processes:
        return "No business processes detected."
    
    summary = "DETECTED BUSINESS PROCESSES:\n"
    
    for process in processes:
        summary += f"\n{process['name']} Process:\n"
        
        # Source and confidence
        summary += f"- Source: {process.get('source', 'unknown')}\n"
        summary += f"- Confidence: {process.get('confidence', 0.5):.0%}\n"
        summary += f"- Complexity: {process['complexity']}\n"
        
        # Add source-specific details
        if process.get('controller'):
            summary += f"- Controller: {process['controller']}\n"
        if process.get('total_actions'):
            summary += f"- Total Actions: {process['total_actions']}\n"
        if process.get('file'):
            summary += f"- File: {process['file']}\n"
        if process.get('tables_involved'):
            summary += f"- Tables Involved: {', '.join(process['tables_involved'])}\n"
        if process.get('has_transaction'):
            summary += f"- Has Transaction: Yes\n"
        
        # CRUD operations (only for controller-based processes)
        crud = process.get('crud_operations', {})
        if crud:
            for operation, actions in crud.items():
                if actions:
                    summary += f"- {operation}: {', '.join(actions)}\n"
        
        # Workflow steps
        workflow_steps = process.get('workflow_steps', [])
        if workflow_steps:
            summary += "- Workflow Steps:\n"
            for step in workflow_steps:
                step_name = step.get('step') or step.get('type', 'Unknown')
                step_type = step.get('type', 'process')
                
                # Safely get roles
                roles = step.get('roles', [])
                if roles and isinstance(roles, list) and roles[0]:
                    roles_text = f" (Roles: {', '.join(str(r) for r in roles if r)})"
                else:
                    roles_text = ""
                
                summary += f"  * {step_name} ({step_type}){roles_text}\n"
    
    return summary


def format_power_platform_mapping(mapping: Dict[str, Any]) -> str:
    """Format Power Platform mapping recommendations"""
    summary = "POWER PLATFORM MAPPING:\n"
    
    # Dataverse tables
    tables = mapping.get('dataverse_tables', [])
    if tables:
        summary += "\nDataverse Tables:\n"
        for table in tables:
            summary += f"- {table.get('legacy_entity', 'Unknown')} → {table.get('suggested_table_name', 'Unknown')}\n"
            summary += f"  Display Name: {table.get('display_name', 'N/A')}\n"
            summary += f"  Schema: {table.get('schema', 'dbo')}\n"
            summary += f"  Columns: {len(table.get('columns', []))}\n"
            
            # Handle sources (plural) from enhanced parser
            sources = table.get('sources', [])
            if sources:
                summary += f"  Sources: {', '.join(sources)}\n"
            
            # Show confidence score
            confidence = table.get('confidence', 0)
            summary += f"  Confidence: {confidence:.0%}\n"
            
            # Flag if needs review
            if table.get('needs_review'):
                summary += f"  ⚠️ Needs Manual Review\n"
    
    # Power Apps screens  
    screens = mapping.get('power_apps_screens', [])
    if screens:
        summary += "\nPower Apps Screens:\n"
        for screen in screens:
            summary += f"- {screen.get('legacy_view', 'Unknown')} → {screen.get('screen_type', 'General')}\n"
            
            # Handle optional fields safely
            if screen.get('fields'):
                summary += f"  Fields: {len(screen['fields'])}\n"
            
            if screen.get('model'):
                summary += f"  Model: {screen['model']}\n"
            
            if screen.get('source_file'):
                summary += f"  Source: {screen['source_file']}\n"
    
    # Power Automate flows
    flows = mapping.get('power_automate_flows', [])
    if flows:
        summary += "\nPower Automate Flows:\n"
        for flow in flows:
            summary += f"- {flow.get('name', 'Unknown Flow')}\n"
            summary += f"  Trigger: {flow.get('trigger', 'Unknown')}\n"
            summary += f"  Steps: {len(flow.get('steps', []))}\n"
            
            if flow.get('business_process'):
                summary += f"  Process: {flow['business_process']}\n"
    
    return summary


def _make_context_snippets(results: List[Dict[str, Any]], max_chars: int = 12000) -> str:
    """Enhanced context snippets with metrics information"""
    parts, total = [], 0
    for r in results or []:
        tag = f"[{r.get('meta',{}).get('path','?')}]"
        
        # Add metrics if available
        metrics_info = ""
        if 'metrics' in r:
            m = r['metrics']
            metrics_info = f" (LOC: {m.get('lines', '?')}, Complexity: {m.get('complexity', '?')})"
        
        piece = f"{tag}{metrics_info}\n{(r['text'] or '').strip()}\n"
        if total + len(piece) > max_chars: 
            break
        parts.append(piece)
        total += len(piece)
    return "\n---\n".join(parts)


# ============================================================
# ✅ NEW: VALIDATION & CONFIDENCE SCORING
# ============================================================

def validate_and_score_extraction(
    parsed_data: Dict[str, Any],
    sample_code_snippets: List[str]
) -> Dict[str, Any]:
    """
    Use GPT-4 to validate extracted data against actual code.
    Returns confidence scores and flags items needing manual review.
    
    Args:
        parsed_data: Dict with power_platform_mapping, business_processes
        sample_code_snippets: List of code strings to validate against
        
    Returns:
        {
            "entities": {"completeness": 0.X, "accuracy": 0.X, "confidence": 0.X, "missing": [...]},
            "processes": {"completeness": 0.X, "accuracy": 0.X, "confidence": 0.X, "missing": [...]},
            "screens": {"completeness": 0.X, "accuracy": 0.X, "confidence": 0.X, "missing": [...]},
            "overall_confidence": 0.X,
            "needs_manual_review": [...]
        }
    """
    
    # Prepare summaries for validation
    entities = parsed_data.get('power_platform_mapping', {}).get('dataverse_tables', [])
    processes = parsed_data.get('business_processes', [])
    screens = parsed_data.get('power_platform_mapping', {}).get('power_apps_screens', [])
    
    entities_summary = "\n".join([
        f"- {e['legacy_entity']}: {len(e.get('columns', []))} columns (confidence: {e.get('confidence', 0):.0%})"
        for e in entities[:10]  # Limit for token efficiency
    ])
    
    processes_summary = "\n".join([
        f"- {p['name']}: {p.get('complexity', 'Unknown')} complexity from {p.get('source', 'unknown')}"
        for p in processes[:10]
    ])
    
    screens_summary = "\n".join([
        f"- {s['legacy_view']}: {s.get('screen_type', 'Unknown')} with {len(s.get('fields', []))} fields"
        for s in screens[:10]
    ])
    
    # Sample code for cross-validation (limit to 2000 chars)
    code_samples = "\n---\n".join([
        snippet[:400] + "..." if len(snippet) > 400 else snippet
        for snippet in sample_code_snippets[:5]
    ])
    
    validation_prompt = f"""You are validating data extraction accuracy from .NET legacy code.

EXTRACTED DATA SUMMARY:

Entities ({len(entities)} total):
{entities_summary}

Business Processes ({len(processes)} total):
{processes_summary}

Screens ({len(screens)} total):
{screens_summary}

SAMPLE CODE TO VALIDATE AGAINST:
{code_samples}

TASK: Assess extraction quality for each category:

1. **Completeness**: Did we capture most items visible in code? (0.0-1.0 score)
2. **Accuracy**: Are extracted items correctly identified? (0.0-1.0 score)
3. **Confidence**: Overall confidence in this category (0.0-1.0 score)
4. **Missing**: List any obvious items we missed (array of strings)

Return ONLY valid JSON:
{{
    "entities": {{
        "completeness": 0.85,
        "accuracy": 0.90,
        "confidence": 0.88,
        "missing": ["PossibleTable1", "PossibleTable2"]
    }},
    "processes": {{
        "completeness": 0.75,
        "accuracy": 0.80,
        "confidence": 0.78,
        "missing": ["ApprovalWorkflow"]
    }},
    "screens": {{
        "completeness": 0.80,
        "accuracy": 0.85,
        "confidence": 0.82,
        "missing": []
    }},
    "overall_confidence": 0.83,
    "needs_manual_review": [
        "Entity 'CustomerAddress' has low confidence (45%)",
        "Process 'Invoice Processing' missing workflow details"
    ]
}}

Be realistic - if code samples don't show certain items, don't penalize. Focus on what's visible.
"""
    
    messages = [
        {"role": "system", "content": "You are a code extraction validation expert. Be precise and realistic with scores."},
        {"role": "user", "content": validation_prompt}
    ]
    
    try:
        response = _chat(messages, temperature=0.0)
        
        # Parse JSON response
        # Try direct parse
        try:
            validation_results = json.loads(response)
        except json.JSONDecodeError:
            # Try extracting JSON from markdown code blocks
            json_match = re.search(r'```(?:json)?\s*(.*?)\s*```', response, re.DOTALL)
            if json_match:
                validation_results = json.loads(json_match.group(1))
            else:
                # Last resort: extract just the JSON object
                json_match = re.search(r'\{.*\}', response, re.DOTALL)
                if json_match:
                    validation_results = json.loads(json_match.group(0))
                else:
                    raise ValueError("Could not extract JSON from response")
        
        # Ensure all required keys exist with defaults
        default_category = {"completeness": 0.5, "accuracy": 0.5, "confidence": 0.5, "missing": []}
        
        validation_results.setdefault("entities", default_category.copy())
        validation_results.setdefault("processes", default_category.copy())
        validation_results.setdefault("screens", default_category.copy())
        validation_results.setdefault("overall_confidence", 0.5)
        validation_results.setdefault("needs_manual_review", ["Validation incomplete - review manually"])
        
        return validation_results
        
    except Exception as e:
        print(f"⚠️  Validation failed: {str(e)}")
        # Return safe defaults if validation fails
        return {
            "entities": {"completeness": 0.6, "accuracy": 0.6, "confidence": 0.6, "missing": []},
            "processes": {"completeness": 0.6, "accuracy": 0.6, "confidence": 0.6, "missing": []},
            "screens": {"completeness": 0.6, "accuracy": 0.6, "confidence": 0.6, "missing": []},
            "overall_confidence": 0.6,
            "needs_manual_review": [f"Validation error: {str(e)}", "Manual review recommended"],
            "validation_error": str(e)
        }


# ============================================================
# ✅ NEW: SEMANTIC BUSINESS RULES EXTRACTION
# ============================================================

def extract_business_rules_from_code(
    code_snippets: List[str],
    nodes: List[Dict[str, Any]]
) -> List[Dict[str, Any]]:
    """
    Extract implicit business rules from conditional logic in code.
    Captures rules buried in if/else statements that aren't in attributes.
    
    Args:
        code_snippets: List of code strings containing business logic
        nodes: Parsed code nodes for context
        
    Returns:
        List of business rules with implementation guidance
    """
    
    # Filter code snippets that likely contain business rules
    rule_candidates = []
    for snippet in code_snippets:
        # Look for conditional statements
        if any(keyword in snippet.lower() for keyword in ['if (', 'if(', 'else if', 'switch', 'case ', '? ', '&&', '||']):
            rule_candidates.append(snippet[:1500])  # Limit size
    
    if not rule_candidates:
        return []
    
    # Take top 10 most promising snippets
    code_sample = "\n\n---CODE SAMPLE---\n\n".join(rule_candidates[:10])
    
    prompt = f"""Extract ALL business rules from this code. Focus on:
- Conditional logic (if/else, switch)
- Validation rules
- Authorization checks
- Approval thresholds
- Status transitions
- Calculations and formulas

CODE SAMPLES:
{code_sample}

For EACH business rule found, return:
{{
    "rule_id": "BR-001",
    "rule_description": "Orders over $5,000 require Director approval",
    "condition_logic": "order.Amount > 5000 AND user.Role != 'Director'",
    "action": "Return Unauthorized or route to approval",
    "field_involved": ["Order.Amount", "User.Role"],
    "dataverse_implementation": "Business Rule: If Amount > 5000, set ApprovalRequired = Yes",
    "power_automate_condition": "If Order Amount is greater than 5000, then start approval flow",
    "priority": "HIGH|MEDIUM|LOW",
    "source": "OrderController.ApproveOrder method"
}}

Return valid JSON array of rules. Extract EVERY rule you find, even simple ones.
CRITICAL: Return ONLY the JSON array, no markdown formatting.
"""
    
    messages = [
        {"role": "system", "content": "You are an expert at extracting business rules from code. Be thorough - extract every rule you find."},
        {"role": "user", "content": prompt}
    ]
    
    try:
        response = _chat(messages, temperature=0.1)
        
        # Try direct parse
        try:
            rules = json.loads(response)
        except json.JSONDecodeError:
            # Try extracting JSON array from markdown
            json_match = re.search(r'```(?:json)?\s*(\[.*?\])\s*```', response, re.DOTALL)
            if json_match:
                rules = json.loads(json_match.group(1))
            else:
                # Try finding array directly
                json_match = re.search(r'\[.*\]', response, re.DOTALL)
                if json_match:
                    rules = json.loads(json_match.group(0))
                else:
                    print("⚠️  Could not parse business rules JSON")
                    return []
        
        # Ensure rules is a list
        if not isinstance(rules, list):
            rules = [rules] if isinstance(rules, dict) else []
        
        print(f"✅ Extracted {len(rules)} business rules from code")
        return rules
        
    except Exception as e:
        print(f"⚠️  Business rules extraction failed: {str(e)}")
        return []


# ============================================================
# ✅ NEW: API INTEGRATION CONTRACT EXTRACTION
# ============================================================

def extract_api_integration_contracts(
    code_snippets: List[str],
    nodes: List[Dict[str, Any]]
) -> List[Dict[str, Any]]:
    """
    Reverse engineer API contracts from HttpClient usage in code.
    Extracts endpoints, methods, auth, request/response schemas, error handling.
    
    Args:
        code_snippets: Code containing HTTP calls
        nodes: Parsed nodes for context
        
    Returns:
        List of API integration contracts
    """
    
    # Filter snippets with HTTP calls
    http_candidates = []
    for snippet in code_snippets:
        if any(keyword in snippet for keyword in ['HttpClient', 'RestSharp', 'HttpPost', 'HttpGet', 'WebClient', 'fetch(', 'axios']):
            http_candidates.append(snippet[:2000])
    
    if not http_candidates:
        return []
    
    code_sample = "\n\n---HTTP CALL SAMPLE---\n\n".join(http_candidates[:8])
    
    prompt = f"""Analyze these HTTP API calls and extract integration contracts.

CODE WITH API CALLS:
{code_sample}

For EACH distinct API integration, extract:
{{
    "integration_name": "ERP Order Sync",
    "endpoint": "https://api.erp.com/orders",
    "method": "POST",
    "authentication": {{
        "type": "OAuth 2.0 / Bearer Token / API Key / Basic",
        "token_source": "Configuration / Azure KeyVault / Hardcoded",
        "details": "client_credentials flow"
    }},
    "request_schema": {{
        "orderNumber": "string",
        "totalAmount": "decimal",
        "items": "array"
    }},
    "response_schema": {{
        "orderId": "string",
        "status": "string"
    }},
    "error_handling": {{
        "400": "Invalid request - log error and notify admin",
        "401": "Authentication failed - refresh token and retry",
        "500": "Server error - retry 3 times with exponential backoff"
    }},
    "retry_logic": "3 attempts with 1s, 2s, 4s delays",
    "timeout": "30 seconds",
    "power_automate_connector": {{
        "connector_type": "HTTP / Custom Connector",
        "authentication_config": "OAuth 2.0 with client credentials",
        "error_handling_steps": "Scope + Configure run after + Send notification"
    }},
    "source_file": "OrderService.cs"
}}

Return valid JSON array. Extract EVERY API call you find.
CRITICAL: Return ONLY the JSON array, no markdown.
"""
    
    messages = [
        {"role": "system", "content": "You are an expert at reverse-engineering API contracts from code. Extract complete integration specs."},
        {"role": "user", "content": prompt}
    ]
    
    try:
        response = _chat(messages, temperature=0.1)
        
        # Try direct parse
        try:
            integrations = json.loads(response)
        except json.JSONDecodeError:
            # Try extracting from markdown
            json_match = re.search(r'```(?:json)?\s*(\[.*?\])\s*```', response, re.DOTALL)
            if json_match:
                integrations = json.loads(json_match.group(1))
            else:
                json_match = re.search(r'\[.*\]', response, re.DOTALL)
                if json_match:
                    integrations = json.loads(json_match.group(0))
                else:
                    print("⚠️  Could not parse integrations JSON")
                    return []
        
        if not isinstance(integrations, list):
            integrations = [integrations] if isinstance(integrations, dict) else []
        
        print(f"✅ Extracted {len(integrations)} API integrations")
        return integrations
        
    except Exception as e:
        print(f"⚠️  Integration extraction failed: {str(e)}")
        return []


# ============================================================
# MAIN BRD GENERATION (WITH ALL ENHANCEMENTS INTEGRATED)
# ============================================================

def generate_brd(retrieved: List[Dict[str, Any]], 
                 nodes: List[Dict[str, Any]], 
                 metrics: Dict[str, Any] = None,
                 business_processes: List[Dict[str, Any]] = None,
                 power_platform_mapping: Dict[str, Any] = None) -> str:
    """Generate comprehensive BRD with metrics and Power Platform focus"""
    try:
        graph_summary = summarize_graph_enhanced(nodes)
        context = _make_context_snippets(retrieved, max_chars=15000)
        
        # Format additional context
        metrics_summary = format_metrics_summary(metrics or {})
        processes_summary = format_business_processes(business_processes or [])
        mapping_summary = format_power_platform_mapping(power_platform_mapping or {})
        
        messages = [
            {"role": "system", "content": BRD_SYSTEM_PROMPT},
            {"role": "user", "content": f"""
            GRAPH:
            {graph_summary}

            {metrics_summary}

            {processes_summary}

            {mapping_summary}

            CONTEXT SNIPPETS:
            {context}
            """}
        ]
    except Exception as e:
        print(f"⚠️  Error preparing BRD context: {str(e.with_traceback(None))}")
        return "Error generating BRD."
    # Generate BRD content
    brd_content = _chat(messages, temperature=0.1)
    # return brd_content
    try:
        # ✅ NEW: Extract business rules from code
        print("Extracting business rules from code...")
        sample_code = [r["text"] for r in retrieved[:10]]  # Get more samples for rules
        business_rules = extract_business_rules_from_code(sample_code, nodes)
        
        # Add business rules section to BRD
        if business_rules:
            rules_section = f"""

    ---

    ## 📋 Business Rules Extracted from Code

    **Total Rules Found:** {len(business_rules)}

    """
            for rule in business_rules:
                priority_emoji = "🔴" if rule.get('priority') == 'HIGH' else "🟡" if rule.get('priority') == 'MEDIUM' else "🟢"
                
                rules_section += f"""### {rule.get('rule_id', 'BR-???')}: {rule.get('rule_description', 'Unknown Rule')} {priority_emoji}

    **Condition:** `{rule.get('condition_logic', 'N/A')}`  
    **Action:** {rule.get('action', 'N/A')}  
    **Fields Involved:** {', '.join(rule.get('field_involved', []))}  

    **Power Platform Implementation:**
    - **Dataverse:** {rule.get('dataverse_implementation', 'N/A')}
    - **Power Automate:** {rule.get('power_automate_condition', 'N/A')}

    **Source:** {rule.get('source', 'Unknown')}

    ---

    """
            
            brd_content += rules_section
        
        # ✅ NEW: Extract API integration contracts
        print("Extracting API integration contracts...")
        api_integrations = extract_api_integration_contracts(sample_code, nodes)
        
        if api_integrations:
            integrations_section = f"""

    ---

    ## 🔌 API Integration Contracts

    **Total Integrations Found:** {len(api_integrations)}

    """
            for integration in api_integrations:
                integrations_section += f"""### {integration.get('integration_name', 'Unknown Integration')}

    **Endpoint:** `{integration.get('method', 'GET')} {integration.get('endpoint', 'N/A')}`  
    **Authentication:** {integration.get('authentication', {}).get('type', 'Unknown')}  
    **Source:** {integration.get('source_file', 'Unknown')}

    #### Request Schema
    ```json
    {json.dumps(integration.get('request_schema', {}), indent=2)}
    ```

    #### Response Schema
    ```json
    {json.dumps(integration.get('response_schema', {}), indent=2)}
    ```

    #### Error Handling
    """
                
                error_handling = integration.get('error_handling', {})
                if error_handling:
                    for status_code, handling in error_handling.items():
                        integrations_section += f"- **{status_code}:** {handling}\n"
                else:
                    integrations_section += "- No specific error handling detected\n"
                
                integrations_section += f"""
    **Retry Logic:** {integration.get('retry_logic', 'None detected')}  
    **Timeout:** {integration.get('timeout', 'Not specified')}

    #### Power Automate Implementation
    - **Connector Type:** {integration.get('power_automate_connector', {}).get('connector_type', 'HTTP')}
    - **Authentication:** {integration.get('power_automate_connector', {}).get('authentication_config', 'Configure manually')}
    - **Error Handling:** {integration.get('power_automate_connector', {}).get('error_handling_steps', 'Add error scopes')}

    ---

    """
            
            brd_content += integrations_section
        
        # ✅ NEW: Add validation report
        print("Running extraction validation...")
        
        validation = validate_and_score_extraction(
            {"power_platform_mapping": power_platform_mapping,
                "business_processes": business_processes
            },
            sample_code
        )
        
        # Append validation section to BRD
        validation_section = f"""

    ---

    ## 🔍 Extraction Quality & Confidence Report

    **Overall Confidence Score:** {validation.get('overall_confidence', 0):.0%}

    | Category | Completeness | Accuracy | Confidence | Status |
    |----------|--------------|----------|------------|--------|
    | **Entities** | {validation.get('entities', {}).get('completeness', 0):.0%} | {validation.get('entities', {}).get('accuracy', 0):.0%} | {validation.get('entities', {}).get('confidence', 0):.0%} | {'✅ Good' if validation.get('entities', {}).get('confidence', 0) >= 0.7 else '⚠️ Review'} |
    | **Processes** | {validation.get('processes', {}).get('completeness', 0):.0%} | {validation.get('processes', {}).get('accuracy', 0):.0%} | {validation.get('processes', {}).get('confidence', 0):.0%} | {'✅ Good' if validation.get('processes', {}).get('confidence', 0) >= 0.7 else '⚠️ Review'} |
    | **Screens** | {validation.get('screens', {}).get('completeness', 0):.0%} | {validation.get('screens', {}).get('accuracy', 0):.0%} | {validation.get('screens', {}).get('confidence', 0):.0%} | {'✅ Good' if validation.get('screens', {}).get('confidence', 0) >= 0.7 else '⚠️ Review'} |

    ### ⚠️ Items Requiring Manual Review

    """
        
        needs_review = validation.get('needs_manual_review', [])
        if needs_review:
            for item in needs_review:
                validation_section += f"- {item}\n"
        else:
            validation_section += "- No items flagged for review\n"
        
        validation_section += "\n### 🔎 Potentially Missed Items\n\n"
        
        any_missing = False
        for category in ['entities', 'processes', 'screens']:
            missing = validation.get(category, {}).get('missing', [])
            if missing:
                validation_section += f"**{category.title()}:**\n"
                for item in missing:
                    validation_section += f"- {item}\n"
                any_missing = True
        
        if not any_missing:
            validation_section += "*No missing items detected in validation sample*\n"
        
        validation_section += "\n**Note:** This validation is based on a sample of the codebase. " \
                            "Items flagged for review or with confidence < 70% should be manually verified against the full source code.\n"
        return brd_content + validation_section
    except Exception as e:
        print(f"⚠️  Error during validation integration: {str(e.with_traceback(None))}")
        validation_section = "\n\n---\n\n## ⚠️ Extraction Validation Error\nAn error occurred during extraction validation. Please review the extracted data manually.\n"
    


# ============================================================
# OTHER GENERATION FUNCTIONS
# ============================================================

def generate_complexity_analysis(metrics: Dict[str, Any], nodes: List[Dict[str, Any]]) -> str:
    """Generate detailed complexity analysis report"""
    
    metrics_summary = format_metrics_summary(metrics)
    graph_summary = summarize_graph_enhanced(nodes)
    
    messages = [
        {"role": "system", "content": COMPLEXITY_ANALYSIS_PROMPT},
        {"role": "user", "content": f"""
{metrics_summary}

{graph_summary}

Provide detailed analysis of code complexity and migration recommendations.
"""}
    ]
    
    return _chat(messages, temperature=0.0)


def generate_business_process_flows(business_processes: List[Dict[str, Any]], 
                                   nodes: List[Dict[str, Any]]) -> str:
    """Generate detailed business process flow documentation"""
    
    processes_summary = format_business_processes(business_processes)
    graph_summary = summarize_graph_enhanced(nodes)
    
    messages = [
        {"role": "system", "content": BUSINESS_PROCESS_FLOW_PROMPT},
        {"role": "user", "content": f"""
{processes_summary}

{graph_summary}

Generate detailed business process flow documentation for Power Platform migration.
"""}
    ]
    
    return _chat(messages, temperature=0.0)


def generate_power_platform_detailed_mapping(power_platform_mapping: Dict[str, Any], 
                                           nodes: List[Dict[str, Any]],
                                           business_processes: List[Dict[str, Any]]) -> str:
    """Generate comprehensive Power Platform component mapping"""
    
    mapping_summary = format_power_platform_mapping(power_platform_mapping)
    processes_summary = format_business_processes(business_processes)
    graph_summary = summarize_graph_enhanced(nodes)
    
    messages = [
        {"role": "system", "content": POWER_PLATFORM_MAPPING_PROMPT},
        {"role": "user", "content": f"""
{mapping_summary}

{processes_summary}

{graph_summary}

Generate comprehensive Power Platform component mapping and migration strategy.
"""}
    ]
    
    return _chat(messages, temperature=0.0)


def generate_user_stories(business_processes: List[Dict[str, Any]], 
                         nodes: List[Dict[str, Any]],
                         power_platform_mapping: Dict[str, Any], *args, **kwargs) -> str:
    """Generate user stories for Power Platform development"""
    
    processes_summary = format_business_processes(business_processes or [])
    mapping_summary = format_power_platform_mapping(power_platform_mapping or {})
    graph_summary = summarize_graph_enhanced(nodes or [])
    
    # Extract personas from nodes (from authorization attributes)
    personas = set()
    for node in nodes:
        if node.get("kind") in ["mvc_controller", "mvc_action"]:
            roles = node.get("props", {}).get("roles", [])
            if roles and isinstance(roles, list):
                for role in roles:
                    if role and isinstance(role, str):
                        personas.update(r.strip() for r in role.split(',') if r.strip())
    
    # Also extract from business processes
    for process in business_processes or []:
        for step in process.get('workflow_steps', []):
            roles = step.get('roles', [])
            if roles and isinstance(roles, list):
                for role in roles:
                    if role and isinstance(role, str):
                        personas.add(role.strip())

    personas_text = f"IDENTIFIED PERSONAS: {', '.join(personas)}" if personas else "No explicit personas found in authorization attributes."
    
    messages = [
        {"role": "system", "content": USER_STORY_GENERATION_PROMPT},
        {"role": "user", "content": f"""
{personas_text}

{processes_summary}

{mapping_summary}

{graph_summary}

Generate comprehensive user stories for Power Platform development teams.
"""}
    ]
    
    return _chat(messages, temperature=0.1)


def answer_question_enhanced(question: str, 
                           retrieved: List[Dict[str, Any]], 
                           nodes: List[Dict[str, Any]],
                           metrics: Dict[str, Any] = None,
                           business_processes: List[Dict[str, Any]] = None) -> str:
    """Enhanced Q&A with metrics and business process context"""
    
    graph_summary = summarize_graph_enhanced(nodes)
    context = _make_context_snippets(retrieved, max_chars=8000)
    
    additional_context = ""
    if metrics:
        additional_context += f"\n{format_metrics_summary(metrics)}"
    if business_processes:
        additional_context += f"\n{format_business_processes(business_processes)}"
    
    messages = [
        {"role": "system", "content": QNA_SYSTEM_PROMPT},
        {"role": "user", "content": f"""
QUESTION:
{question}

GRAPH:
{graph_summary}

{additional_context}

CONTEXT:
{context}
"""}
    ]
    
    return _chat(messages, temperature=0.0)


# ============================================================
# ASYNC DOCUMENT GENERATORS FOR TAB 6
# ============================================================

import asyncio
from datetime import datetime

async def generate_business_flows(retrieved: List[Dict[str, Any]], 
                         nodes: List[Dict[str, Any]], 
                         metrics: Dict[str, Any] = None,
                         business_processes: List[Dict[str, Any]] = None,
                         power_platform_mapping: Dict[str, Any] = None) -> str:
    """Generate comprehensive business process flows documentation"""
    
    graph_summary = summarize_graph_enhanced(nodes)
    context = _make_context_snippets(retrieved, max_chars=15000)
    
    # Format additional context
    metrics_summary = format_metrics_summary(metrics or {})
    processes_summary = format_business_processes(business_processes or [])
    mapping_summary = format_power_platform_mapping(power_platform_mapping or {})
    
    messages = [
        {"role": "system", "content": "Generate detailed business process flow documentation for .NET to Power Platform migration. Create a structured document with process flows, decision points, and Power Platform mappings."},
        {"role": "user", "content": f"""
        GRAPH:
        {graph_summary}

        {metrics_summary}

        {processes_summary}

        {mapping_summary}

        CONTEXT SNIPPETS:
        {context}
        """}
    ]
    
    return _chat(messages, temperature=0.1)


async def generate_tables_analysis(retrieved: List[Dict[str, Any]], 
                         nodes: List[Dict[str, Any]], 
                         metrics: Dict[str, Any] = None,
                         business_processes: List[Dict[str, Any]] = None,
                         power_platform_mapping: Dict[str, Any] = None) -> str:
    """Generate comprehensive database and Dataverse mapping documentation"""
    
    graph_summary = summarize_graph_enhanced(nodes)
    context = _make_context_snippets(retrieved, max_chars=15000)
    
    # Format additional context
    metrics_summary = format_metrics_summary(metrics or {})
    processes_summary = format_business_processes(business_processes or [])
    mapping_summary = format_power_platform_mapping(power_platform_mapping or {})
    
    messages = [
        {"role": "system", "content": "Generate detailed Dataverse table mapping documentation for .NET to Power Platform migration. Create a structured document with table schemas, relationships, and migration strategies."},
        {"role": "user", "content": f"""
        GRAPH:
        {graph_summary}

        {metrics_summary}

        {processes_summary}

        {mapping_summary}

        CONTEXT SNIPPETS:
        {context}
        """}
    ]
    
    return _chat(messages, temperature=0.1)


async def generate_user_journeys(retrieved: List[Dict[str, Any]], 
                         nodes: List[Dict[str, Any]], 
                         metrics: Dict[str, Any] = None,
                         business_processes: List[Dict[str, Any]] = None,
                         power_platform_mapping: Dict[str, Any] = None) -> str:
    """Generate comprehensive user journey documentation"""
    
    graph_summary = summarize_graph_enhanced(nodes)
    context = _make_context_snippets(retrieved, max_chars=15000)
    
    # Format additional context
    metrics_summary = format_metrics_summary(metrics or {})
    processes_summary = format_business_processes(business_processes or [])
    mapping_summary = format_power_platform_mapping(power_platform_mapping or {})
    
    messages = [
        {"role": "system", "content": "Generate detailed user journey documentation for .NET to Power Platform migration. Create step-by-step user flows showing how different personas interact with the system."},
        {"role": "user", "content": f"""
        GRAPH:
        {graph_summary}

        {metrics_summary}

        {processes_summary}

        {mapping_summary}

        CONTEXT SNIPPETS:
        {context}
        """}
    ]
    
    return _chat(messages, temperature=0.1)


async def generate_personas(retrieved: List[Dict[str, Any]], 
                         nodes: List[Dict[str, Any]], 
                         metrics: Dict[str, Any] = None,
                         business_processes: List[Dict[str, Any]] = None,
                         power_platform_mapping: Dict[str, Any] = None) -> str:
    """Generate persona documentation"""
    
    graph_summary = summarize_graph_enhanced(nodes)
    context = _make_context_snippets(retrieved, max_chars=15000)
    
    # Format additional context
    metrics_summary = format_metrics_summary(metrics or {})
    processes_summary = format_business_processes(business_processes or [])
    mapping_summary = format_power_platform_mapping(power_platform_mapping or {})
    
    messages = [
        {"role": "system", "content": "Generate persona documentation for .NET to Power Platform migration. Create detailed user personas with roles, responsibilities, goals, and pain points."},
        {"role": "user", "content": f"""
        GRAPH:
        {graph_summary}

        {metrics_summary}

        {processes_summary}

        {mapping_summary}

        CONTEXT SNIPPETS:
        {context}
        """}
    ]
    
    return _chat(messages, temperature=0.1)


async def generate_user_storie(analysis_data: Dict[str, Any]) -> str:
    """Generate user stories for Power Platform development"""
    
    app_name = analysis_data.get('application_name', 'Application')
    business_processes = analysis_data.get('business_processes', [])
    power_mapping = analysis_data.get('power_platform_mapping', {})
    
    content = f"""# User Stories - {app_name}

## Document Information
- **Generated Date**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
- **Purpose**: Agile user stories for Power Platform development
- **Sprint Planning**: Ready for estimation and sprint allocation

---

## Epic: {app_name} Power Platform Migration

"""
    
    story_id = 1
    
    # Generate stories for each business process
    for process in business_processes:
        content += f"\n### Feature: {process['name']}\n\n"
        
        # Story for main process
        content += f"""#### US-{story_id:03d}: {process['name']} - Main Process

**As a** {', '.join(process.get('workflow_steps', [{}])[0].get('roles', ['User'])) if process.get('workflow_steps') else 'User'}
**I want to** execute the {process['name']} process
**So that** I can complete my business operations efficiently

**Acceptance Criteria:**
"""
        story_id += 1
        
        crud = process.get('crud_operations', {})
        for action in crud.get('read', []):
            content += f"- [ ] User can view {action} data\n"
        
        for action in crud.get('create', []):
            content += f"- [ ] User can create new {action} records\n"
        
        for action in crud.get('update', []):
            content += f"- [ ] User can update existing {action} records\n"
        
        for action in crud.get('delete', []):
            content += f"- [ ] User can delete {action} records\n"
        
        content += f"\n**Story Points**: {process.get('total_actions', 0) * 2}\n"
        content += f"**Priority**: {'High' if process['complexity'] == 'High' else 'Medium'}\n\n"
        content += "---\n"
    
    return content


async def generate_test_cases(retrieved: List[Dict[str, Any]], 
                         nodes: List[Dict[str, Any]], 
                         metrics: Dict[str, Any] = None,
                         business_processes: List[Dict[str, Any]] = None,
                         power_platform_mapping: Dict[str, Any] = None) -> str:
    """Generate functional test cases documentation"""
    
    graph_summary = summarize_graph_enhanced(nodes)
    context = _make_context_snippets(retrieved, max_chars=15000)
    
    # Format additional context
    metrics_summary = format_metrics_summary(metrics or {})
    processes_summary = format_business_processes(business_processes or [])
    mapping_summary = format_power_platform_mapping(power_platform_mapping or {})
    
    messages = [
        {"role": "system", "content": "Generate comprehensive test case documentation for .NET to Power Platform migration. Create functional test scenarios with test steps, expected results, and test data."},
        {"role": "user", "content": f"""
        GRAPH:
        {graph_summary}

        {metrics_summary}

        {processes_summary}

        {mapping_summary}

        CONTEXT SNIPPETS:
        {context}
        """}
    ]
    
    return _chat(messages, temperature=0.1)


async def generate_complete_brd_async(analysis_data: Dict[str, Any]) -> str:
    """Generate complete BRD with all sections (async wrapper for compatibility)"""
    
    # This is a compatibility function - the main generate_brd already does everything
    return generate_brd(
        analysis_data.get('retrieved', []),
        analysis_data.get('nodes', []),
        analysis_data.get('metrics', {}),
        analysis_data.get('business_processes', []),
        analysis_data.get('power_platform_mapping', {})
    )


# ============================================================
# COPILOT PROMPT LIBRARY GENERATOR (OPTIONAL ENHANCEMENT)
# ============================================================

def generate_copilot_prompt_library(
    retrieved: List[Dict[str, Any]],
    nodes: List[Dict[str, Any]],
    metrics: Dict[str, Any],
    business_processes: List[Dict[str, Any]],
    power_platform_mapping: Dict[str, Any]
) -> Dict[str, Any]:
    """
    Generate comprehensive, copy-paste ready Copilot prompts.
    Uses GPT-4 to create natural language instructions for each component.
    """
    
    # Prepare context
    graph_summary = summarize_graph_enhanced(nodes)
    context_snippets = _make_context_snippets(retrieved, max_chars=8000)
    processes_summary = format_business_processes(business_processes or [])
    mapping_summary = format_power_platform_mapping(power_platform_mapping or {})
    
    # Enhanced system prompt for Copilot prompt generation
    system_prompt = """You are an expert at creating natural language prompts for Microsoft Copilot tools in Power Platform.

Your task: Generate ready-to-use prompts that developers can copy-paste directly into:
1. Copilot in Power Apps (canvas apps)
2. Copilot in Power Automate (cloud flows)  
3. Copilot for Dataverse (table creation)

REQUIREMENTS for each prompt:
- Natural conversational language (not technical jargon)
- Complete specifications (no placeholders like "add your fields here")
- Include validation rules and business logic explicitly
- Specify navigation and user interactions clearly
- Add error handling instructions
- Maximum 250 words per prompt (Copilot's optimal length)
- Each prompt must be self-contained and actionable

OUTPUT FORMAT: Valid JSON only with this structure:
{
    "dataverse_prompts": [
        {
            "id": "DV-001",
            "title": "Create Customer table",
            "priority": "HIGH|MEDIUM|LOW",
            "copilot_tool": "Copilot in Dataverse",
            "prompt": "Natural language prompt text...",
            "validation": ["Check 1", "Check 2", "Check 3"],
            "estimated_time": "15-20 min"
        }
    ],
    "power_apps_prompts": [...],
    "power_automate_prompts": [...]
}"""

    user_prompt = f"""Generate Copilot prompts for migrating this .NET application to Power Platform.

ANALYZED CODE STRUCTURE:
{graph_summary}

BUSINESS PROCESSES:
{processes_summary}

POWER PLATFORM MAPPING:
{mapping_summary}

CODE SAMPLES:
{context_snippets[:3000]}

Generate prompts for:
1. **Top 5 Dataverse Tables** (most important entities)
2. **Top 5 Power Apps Screens** (critical user interfaces)
3. **Top 3 Power Automate Flows** (key workflows)

For EACH prompt include:
- "id": unique identifier (e.g., "DV-001", "PA-001", "PAF-001")
- "title": short description (max 50 chars)
- "priority": HIGH (critical path), MEDIUM (important), or LOW (nice-to-have)
- "copilot_tool": which Copilot tool to use
- "prompt": the actual conversational prompt (150-250 words, complete, actionable)
- "validation": array of 3-5 checkpoints to verify it worked
- "estimated_time": realistic time estimate with Copilot (e.g., "15-20 min")

CRITICAL: Each prompt must be complete enough that a developer can paste it into Copilot and get a working component without additional clarification.

Return ONLY valid JSON, no markdown formatting."""

    # Call GPT-4 with structured output
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt}
    ]
    
    response = _chat(messages, temperature=0.2)  # Slightly higher for creativity
    
    # Parse response (handle potential JSON errors)
    try:
        library = json.loads(response)
    except json.JSONDecodeError:
        # Fallback: extract JSON from markdown code blocks
        json_match = re.search(r'```(?:json)?\s*(.*?)\s*```', response, re.DOTALL)
        if json_match:
            try:
                library = json.loads(json_match.group(1))
            except:
                library = {"error": "JSON parse failed", "raw_response": response[:500]}
        else:
            # Last attempt: find JSON object
            json_match = re.search(r'\{.*\}', response, re.DOTALL)
            if json_match:
                try:
                    library = json.loads(json_match.group(0))
                except:
                    library = {"error": "Could not extract JSON"}
            else:
                library = {
                    "dataverse_prompts": [],
                    "power_apps_prompts": [],
                    "power_automate_prompts": [],
                    "error": "Failed to parse GPT response",
                    "raw_response": response[:500]
                }
    
    # Ensure all expected keys exist
    library.setdefault("dataverse_prompts", [])
    library.setdefault("power_apps_prompts", [])
    library.setdefault("power_automate_prompts", [])
    
    # Add metadata
    import pandas as pd
    library["metadata"] = {
        "generated_date": pd.Timestamp.now().isoformat(),
        "total_prompts": sum(
            len(v) for k, v in library.items() 
            if isinstance(v, list) and k.endswith('_prompts')
        ),
        "source_files": len(nodes),
        "business_processes": len(business_processes or []),
        "entities_analyzed": len(power_platform_mapping.get('dataverse_tables', []))
    }
    
    return library


def format_copilot_library_as_markdown(library: Dict[str, Any]) -> str:
    """Convert JSON library to readable Markdown for download"""
    
    md_parts = [
        f"# 🤖 Copilot Prompt Library\n\n",
        f"**Generated:** {library['metadata']['generated_date']}\n",
        f"**Total Prompts:** {library['metadata']['total_prompts']}\n",
        f"**Source Files Analyzed:** {library['metadata']['source_files']}\n",
        f"**Business Processes:** {library['metadata']['business_processes']}\n\n",
        "---\n\n"
    ]
    
    # Dataverse section
    if library.get('dataverse_prompts'):
        md_parts.append("## 📊 Dataverse Table Creation Prompts\n\n")
        md_parts.append("*Copy-paste these into **Copilot in Dataverse***\n\n")
        
        for i, prompt in enumerate(library['dataverse_prompts'], 1):
            priority_emoji = "🔴" if prompt.get('priority') == 'HIGH' else "🟡" if prompt.get('priority') == 'MEDIUM' else "🟢"
            
            md_parts.extend([
                f"### {i}. {prompt.get('title', 'Untitled')} {priority_emoji}\n\n",
                f"**Priority:** {prompt.get('priority', 'MEDIUM')} | ",
                f"**Estimated Time:** {prompt.get('estimated_time', '15-20 min')}\n\n",
                "**📋 Copilot Prompt:**\n\n",
                "```\n",
                prompt.get('prompt', 'N/A'),
                "\n```\n\n",
                "**✅ Validation Checklist:**\n\n"
            ])
            
            for check in prompt.get('validation', []):
                md_parts.append(f"- [ ] {check}\n")
            
            md_parts.append("\n---\n\n")
    
    # Power Apps section
    if library.get('power_apps_prompts'):
        md_parts.append("## 📱 Power Apps Canvas Screen Prompts\n\n")
        md_parts.append("*Copy-paste these into **Copilot in Power Apps***\n\n")
        
        for i, prompt in enumerate(library['power_apps_prompts'], 1):
            priority_emoji = "🔴" if prompt.get('priority') == 'HIGH' else "🟡" if prompt.get('priority') == 'MEDIUM' else "🟢"
            
            md_parts.extend([
                f"### {i}. {prompt.get('title', 'Untitled')} {priority_emoji}\n\n",
                f"**Priority:** {prompt.get('priority', 'MEDIUM')} | ",
                f"**Estimated Time:** {prompt.get('estimated_time', '20-30 min')}\n\n",
                "**📋 Copilot Prompt:**\n\n",
                "```\n",
                prompt.get('prompt', 'N/A'),
                "\n```\n\n",
                "**✅ Validation Checklist:**\n\n"
            ])
            
            for check in prompt.get('validation', []):
                md_parts.append(f"- [ ] {check}\n")
            
            md_parts.append("\n---\n\n")
    
    # Power Automate section
    if library.get('power_automate_prompts'):
        md_parts.append("## ⚡ Power Automate Flow Prompts\n\n")
        md_parts.append("*Copy-paste these into **Copilot in Power Automate***\n\n")
        
        for i, prompt in enumerate(library['power_automate_prompts'], 1):
            priority_emoji = "🔴" if prompt.get('priority') == 'HIGH' else "🟡" if prompt.get('priority') == 'MEDIUM' else "🟢"
            
            md_parts.extend([
                f"### {i}. {prompt.get('title', 'Untitled')} {priority_emoji}\n\n",
                f"**Priority:** {prompt.get('priority', 'MEDIUM')} | ",
                f"**Estimated Time:** {prompt.get('estimated_time', '30-40 min')}\n\n",
                "**📋 Copilot Prompt:**\n\n",
                "```\n",
                prompt.get('prompt', 'N/A'),
                "\n```\n\n",
                "**✅ Validation Checklist:**\n\n"
            ])
            
            for check in prompt.get('validation', []):
                md_parts.append(f"- [ ] {check}\n")
            
            md_parts.append("\n---\n\n")
    
    # Quick start guide
    md_parts.extend([
        "## 🚀 Quick Start Guide\n\n",
        "### How to Use These Prompts:\n\n",
        "1. **Copy** the prompt text from inside the code block above\n",
        "2. **Open** the corresponding Copilot tool in Power Platform\n",
        "3. **Paste** the entire prompt into the Copilot input box\n",
        "4. **Press Enter** and wait for Copilot to generate the component\n",
        "5. **Review** the generated result and make minor adjustments if needed\n",
        "6. **Validate** using the checklist provided with each prompt\n",
        "7. **Test** the component with sample data\n\n",
        "### 💡 Tips for Best Results:\n\n",
        "- **Follow the order:** Create Dataverse tables first, then apps, then flows\n",
        "- **Validate each step:** Check off the validation items before proceeding\n",
        "- **Customize as needed:** These prompts are starting points\n",
        "- **Test incrementally:** Don't build everything at once\n\n",
        "---\n\n",
        "*Generated by .NET to Power Platform Migration Assistant*\n"
    ])
    
    return ''.join(md_parts)



