from utils import chunk_document_safe
import json, re, time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
from datetime import datetime

def validate_uploaded_brd(brd_text: str, seed_results: list) -> dict:
    """
    ✅ Enhanced BRD validation with dynamic section detection and detailed failure reasons.
    Validates uploaded BRD (supports 100+ pages) against repository insights.
    Uses shared _chat() for Azure GPT-4.1-mini.
    """

    # --- Context summary (limit safely) ---
    context_summary = "\n\n".join([r.get("text", "")[:2000] for r in seed_results[:10]])

    # --- Chunk entire BRD (~9k tokens each ≈ 20–25 pages) ---
    brd_chunks = chunk_document_safe(brd_text, max_tokens=9000)
    total_chunks = len(brd_chunks)
    print(f"🧾 Validating BRD ({total_chunks} chunks, full coverage)...")

    all_results = []

    for i, chunk in enumerate(brd_chunks, start=1):
        prompt = f"""
You are a STRICT BRD validator with expertise in software requirements analysis.

**YOUR CRITICAL TASK:**
Compare this BRD content chunk ({i}/{total_chunks}) with the PROVIDED REPOSITORY CONTEXT.
You MUST validate if the BRD describes the SAME application as the repository context.

**STRICT VALIDATION PROTOCOL:**

1. **VERIFY APPLICATION MATCH:**
   - Does the BRD application name/domain match the repository context?
   - Are the core features/modules mentioned in BRD present in the repository?
   - Do technologies, frameworks, and architectures align?
   - If the BRD describes a DIFFERENT application → FAIL immediately

2. **CROSS-REFERENCE REQUIREMENTS:**
   For each section in the BRD chunk, you MUST:
   - Find SPECIFIC evidence in the repository context that supports it
   - If NO matching evidence exists in repository → Status = FAIL
   - If requirements contradict repository details → Status = FAIL
   - If requirements mention features NOT in repository → Status = FAIL

3. **MANDATORY EVIDENCE-BASED VALIDATION:**
   - PASS only if you can cite specific repository context that confirms the BRD claim
   - FAIL if BRD mentions entities, features, modules, APIs, or workflows NOT found in context
   - FAIL if technology stack doesn't match (e.g., BRD says React but repo uses Angular)
   - FAIL if data models, database tables, or schemas don't align

**RESPOND IN JSON:**
{{
  "status": "valid" or "invalid",
  "summary": "State clearly if BRD matches repository or describes different application",
  "application_match": "yes" or "no" or "uncertain",
  "sections": [
    {{
      "name": "Section name from BRD",
      "status": "pass" or "fail",
      "reason": "MANDATORY: Cite specific repository evidence if PASS, or explain mismatch if FAIL",
      "confidence": "high" or "medium" or "low",
      "repository_evidence": "Quote relevant repository context that supports/contradicts this section"
    }}
  ],
  "detected_sections_count": <number>,
  "critical_issues": ["List mismatches: wrong app, missing features, tech conflicts, etc."],
  "misalignment_score": <0-100, where 0=perfect match, 100=completely different app>
}}

**VALIDATION EXAMPLES:**

✅ PASS Example:
- BRD: "User authentication via JWT tokens"
- Repository: "JWT authentication middleware in auth.js, login endpoint /api/auth/login"
- Status: PASS - Evidence found

❌ FAIL Example:
- BRD: "E-commerce shopping cart with payment gateway"
- Repository: "Hospital management system with patient records and appointments"
- Status: FAIL - Completely different application domain

❌ FAIL Example:
- BRD: "React frontend with Redux state management"
- Repository: "Angular application using NgRx"
- Status: FAIL - Technology mismatch

**BE STRICT:** Default to FAIL unless you have clear repository evidence.

--- REPOSITORY CONTEXT (THIS IS THE SOURCE OF TRUTH) ---
{context_summary}

--- BRD CHUNK TO VALIDATE ({i}/{total_chunks}) ---
{chunk[:8500]}

Validate strictly: Does this BRD chunk describe the SAME system as the repository context?
"""

        # --- Retry logic using _chat() ---
        for attempt in range(3):
            try:
                response_text = _chat([
                    {"role": "system", "content": "You are an enterprise BRD validator. Respond only in valid JSON format."},
                    {"role": "user", "content": prompt}
                ], temperature=0.1)

                # Extract JSON from response
                json_match = re.search(r"\{.*\}", response_text, re.DOTALL)
                if not json_match:
                    raise ValueError("No JSON object found in model output")

                parsed = json.loads(json_match.group(0))

                # Normalize structure
                if "sections" not in parsed or not isinstance(parsed["sections"], list):
                    parsed["sections"] = []
                if "critical_issues" not in parsed:
                    parsed["critical_issues"] = []
                if "detected_sections_count" not in parsed:
                    parsed["detected_sections_count"] = len(parsed["sections"])
                if "application_match" not in parsed:
                    parsed["application_match"] = "uncertain"
                if "misalignment_score" not in parsed:
                    parsed["misalignment_score"] = 50  # Default to medium risk

                parsed["chunk_id"] = i
                all_results.append(parsed)
                
                # Log alignment score for monitoring
                score = parsed.get("misalignment_score", 50)
                match = parsed.get("application_match", "uncertain")
                print(f"✓ Chunk {i}/{total_chunks}: {parsed['detected_sections_count']} sections, Match={match}, Misalignment={score}%")
                break
                
            except Exception as e:
                print(f"⚠️ Chunk {i}/{total_chunks} failed (attempt {attempt+1}): {e}")
                time.sleep(2 + attempt)
                if attempt == 2:
                    # Fallback entry when the model fails 3 times
                    all_results.append({
                        "chunk_id": i,
                        "status": "invalid",
                        "summary": f"Chunk {i} failed validation due to model or parsing errors after 3 attempts.",
                        "sections": [{
                            "name": "Validation Error",
                            "status": "fail",
                            "reason": f"Technical error during validation: {str(e)}. Unable to process this section of the BRD.",
                            "confidence": "low"
                        }],
                        "detected_sections_count": 0,
                        "critical_issues": ["Validation process failure"]
                    })

    if not all_results:
        return {
            "status": "invalid",
            "summary": "No valid responses from validation model.",
            "sections": {},
            "chunks": [],
            "critical_issues": ["Complete validation failure"]
        }

    # --- Aggregate Results with Enhanced Tracking ---
    section_stats = {}
    valid_count = 0
    invalid_count = 0
    all_critical_issues = []
    total_misalignment_score = 0
    application_match_votes = {"yes": 0, "no": 0, "uncertain": 0}

    for r in all_results:
        # Track chunk-level status
        if r.get("status", "invalid").lower() == "valid":
            valid_count += 1
        else:
            invalid_count += 1

        # Track application match consensus
        app_match = r.get("application_match", "uncertain")
        application_match_votes[app_match] = application_match_votes.get(app_match, 0) + 1
        
        # Accumulate misalignment scores
        total_misalignment_score += r.get("misalignment_score", 50)

        # Collect critical issues
        critical = r.get("critical_issues", [])
        if critical:
            all_critical_issues.extend(critical)

        # Aggregate section-level results
        for sec in r.get("sections", []):
            name = (sec.get("name") or "").strip()
            if not name:
                continue
                
            status = (sec.get("status") or "fail").lower()
            reason = (sec.get("reason") or "").strip()
            confidence = sec.get("confidence", "medium")
            evidence = sec.get("repository_evidence", "")

            if name not in section_stats:
                section_stats[name] = {
                    "pass": 0,
                    "fail": 0,
                    "fail_reasons": [],
                    "pass_confirmations": [],
                    "confidence_scores": [],
                    "evidence": []
                }

            section_stats[name]["confidence_scores"].append(confidence)
            if evidence:
                section_stats[name]["evidence"].append(evidence)
            
            if status == "pass":
                section_stats[name]["pass"] += 1
                if reason:
                    section_stats[name]["pass_confirmations"].append(reason)
            else:
                section_stats[name]["fail"] += 1
                if reason:
                    section_stats[name]["fail_reasons"].append(reason)

    # --- Build Final Aggregated Sections with Detailed Reasons ---
    aggregated_sections = {}
    for name, stats in section_stats.items():
        is_pass = stats["pass"] >= stats["fail"]
        status = "pass" if is_pass else "fail"

        # Determine confidence
        confidence_map = {"high": 3, "medium": 2, "low": 1}
        avg_confidence_score = sum(confidence_map.get(c, 2) for c in stats["confidence_scores"]) / len(stats["confidence_scores"])
        if avg_confidence_score >= 2.5:
            confidence = "high"
        elif avg_confidence_score >= 1.5:
            confidence = "medium"
        else:
            confidence = "low"

        # Build detailed reason
        reason = ""
        repository_evidence = ""
        
        if status == "fail":
            # Deduplicate and format failure reasons
            unique_reasons = list(dict.fromkeys(stats["fail_reasons"]))
            if unique_reasons:
                reason = " | ".join(unique_reasons[:5])  # Show up to 5 distinct reasons
            else:
                reason = "Multiple alignment issues detected. No matching evidence found in repository context."
        else:
            # For passing sections, include confirmation
            if stats["pass_confirmations"]:
                reason = stats["pass_confirmations"][0]
        
        # Include repository evidence if available
        if stats["evidence"]:
            repository_evidence = stats["evidence"][0]  # Use first evidence

        aggregated_sections[name] = {
            "status": status,
            "reason": reason,
            "confidence": confidence,
            "pass_count": stats["pass"],
            "fail_count": stats["fail"],
            "repository_evidence": repository_evidence
        }

    # --- Determine Application Match Consensus ---
    avg_misalignment = total_misalignment_score / len(all_results) if all_results else 100
    
    # Determine overall match
    if application_match_votes["no"] > len(all_results) * 0.3:  # If >30% say NO
        application_match = "no"
    elif application_match_votes["yes"] > len(all_results) * 0.5:  # If >50% say YES
        application_match = "yes"
    else:
        application_match = "uncertain"

    # --- Overall Status (STRICTER LOGIC) ---
    # Consider it invalid if:
    # 1. More invalid chunks than valid
    # 2. High misalignment score (>60)
    # 3. Application doesn't match
    # 4. More failed sections than passed
    
    total_sections = len(aggregated_sections)
    failed_sections = sum(1 for s in aggregated_sections.values() if s["status"] == "fail")
    
    overall_status = "valid"
    if invalid_count > valid_count:
        overall_status = "invalid"
    elif avg_misalignment > 60:
        overall_status = "invalid"
    elif application_match == "no":
        overall_status = "invalid"
    elif failed_sections > total_sections * 0.5:  # More than 50% sections failed
        overall_status = "invalid"
    
    # --- Enhanced Summary ---
    summary = (
        f"Application Match: {application_match.upper()}. "
        f"Misalignment Score: {avg_misalignment:.1f}%. "
        f"Validated {total_chunks} chunks covering {total_sections} sections. "
        f"{valid_count} chunks aligned, {invalid_count} misaligned. "
        f"{failed_sections} sections failed validation."
    )

    # Add critical issues to summary if any
    if all_critical_issues:
        unique_critical = list(dict.fromkeys(all_critical_issues))[:3]
        summary += f" CRITICAL: {', '.join(unique_critical)}"

    aggregated = {
        "status": overall_status,
        "summary": summary,
        "sections": aggregated_sections,
        "chunks": all_results,
        "critical_issues": list(dict.fromkeys(all_critical_issues)),
        "statistics": {
            "total_chunks": total_chunks,
            "valid_chunks": valid_count,
            "invalid_chunks": invalid_count,
            "total_sections": total_sections,
            "passed_sections": total_sections - failed_sections,
            "failed_sections": failed_sections
        }
    }

    print(f"✅ BRD Validation Completed: {total_sections} sections across {total_chunks} chunks")
    return aggregated


def generate_brd_validation_docx(result: dict, seed_results: list) -> bytes:
    """
    📄 Generate an enhanced, professional Word report (.docx) for BRD validation.
    Includes dynamic checklist, detailed failure reasons, and visual formatting.
    Returns raw bytes ready for Streamlit download.
    """

    doc = Document()

    # --- Cover Page ---
    title = doc.add_heading("BRD Validation Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(16, 185, 129)  # Green color

    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", style="Body Text").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Validated using Azure GPT-4.1-mini against repository insights", style="Body Text").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # --- Executive Summary ---
    doc.add_heading("📊 Executive Summary", level=1)
    
    status = (result.get("status", "unknown") or "").upper()
    status_text = "✅ VALID" if status == "VALID" else "❌ INVALID"
    status_para = doc.add_paragraph()
    status_run = status_para.add_run(f"Overall Validation Status: {status_text}")
    status_run.bold = True
    status_run.font.size = Pt(12)
    if status == "VALID":
        status_run.font.color.rgb = RGBColor(16, 185, 129)
    else:
        status_run.font.color.rgb = RGBColor(239, 68, 68)

    doc.add_paragraph(result.get("summary", "No summary available."))

    # Statistics
    stats = result.get("statistics", {})
    if stats:
        doc.add_paragraph()
        doc.add_paragraph("📈 Validation Statistics:", style="Heading 3")
        doc.add_paragraph(f"• Total Chunks Analyzed: {stats.get('total_chunks', 0)}")
        doc.add_paragraph(f"• Valid Chunks: {stats.get('valid_chunks', 0)}")
        doc.add_paragraph(f"• Invalid Chunks: {stats.get('invalid_chunks', 0)}")
        doc.add_paragraph(f"• Total Sections Found: {stats.get('total_sections', 0)}")
        doc.add_paragraph(f"• Passed Sections: {stats.get('passed_sections', 0)}")
        doc.add_paragraph(f"• Failed Sections: {stats.get('failed_sections', 0)}")

    # Critical Issues
    critical_issues = result.get("critical_issues", [])
    if critical_issues:
        doc.add_paragraph()
        doc.add_heading("⚠️ Critical Issues", level=2)
        for issue in critical_issues[:10]:  # Limit to 10
            p = doc.add_paragraph(issue, style="List Bullet")
            for run in p.runs:
                run.font.color.rgb = RGBColor(239, 68, 68)

    doc.add_page_break()

    # --- Dynamic Section Checklist ---
    doc.add_heading("✅ BRD Section Validation Checklist", level=1)
    doc.add_paragraph("This checklist is dynamically generated based on sections detected in your BRD.")

    sections = result.get("sections", {})
    if sections:
        # Create table with headers
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Section Name"
        hdr_cells[1].text = "Status"
        hdr_cells[2].text = "Confidence"
        hdr_cells[3].text = "Occurrences"
        hdr_cells[4].text = "Validation Details / Reason"

        # Make headers bold
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # Add section data
        for name, info in sorted(sections.items()):
            row_cells = table.add_row().cells
            
            # Section name
            row_cells[0].text = name
            
            # Status with icon
            status_val = (info.get("status", "") or "").lower()
            status_text = "✅ Pass" if status_val == "pass" else "❌ Fail"
            row_cells[1].text = status_text
            
            # Confidence
            confidence = info.get("confidence", "medium")
            row_cells[2].text = confidence.capitalize()
            
            # Occurrences
            pass_count = info.get("pass_count", 0)
            fail_count = info.get("fail_count", 0)
            row_cells[3].text = f"✓{pass_count} / ✗{fail_count}"
            
            # Reason/Details
            reason = info.get("reason", "") or "No specific details provided."
            row_cells[4].text = reason
            
            # Color coding for failed sections
            if status_val == "fail":
                for paragraph in row_cells[1].paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(239, 68, 68)

        # Summary after table
        doc.add_paragraph()
        passed = sum(1 for s in sections.values() if s.get("status") == "pass")
        failed = sum(1 for s in sections.values() if s.get("status") == "fail")
        doc.add_paragraph(f"Summary: {passed} sections passed, {failed} sections failed out of {len(sections)} total sections detected.")
        
    else:
        doc.add_paragraph("⚠️ No sections were detected in the uploaded BRD. Please verify the document format.")

    doc.add_page_break()

    # --- Repository Context Reference ---
    doc.add_heading("📚 Repository Context Summary", level=1)
    doc.add_paragraph("The following repository insights were used as validation baseline:")
    
    for i, r in enumerate(seed_results[:8], start=1):
        snippet = (r.get("text", "") or "").strip().replace("\n", " ")
        if snippet:
            doc.add_paragraph(f"{i}. {snippet[:600]}{'...' if len(snippet) > 600 else ''}", style="List Number")

    doc.add_page_break()

    # --- Detailed Chunk Analysis ---
    doc.add_heading("🔍 Detailed Chunk-Level Analysis", level=1)
    doc.add_paragraph("Below is a detailed breakdown of validation results for each document chunk:")

    for chunk in result.get("chunks", []):
        chunk_id = chunk.get("chunk_id", "?")
        chunk_status = chunk.get("status", "unknown")
        
        # Chunk header
        heading = doc.add_heading(f"Chunk {chunk_id} - {chunk_status.upper()}", level=2)
        if chunk_status.lower() == "invalid":
            for run in heading.runs:
                run.font.color.rgb = RGBColor(239, 68, 68)
        
        doc.add_paragraph(chunk.get("summary", "No summary available."))
        
        sections_in_chunk = chunk.get("sections", [])
        detected_count = chunk.get("detected_sections_count", len(sections_in_chunk))
        doc.add_paragraph(f"Sections detected in this chunk: {detected_count}")

        if sections_in_chunk:
            # Create mini-table for chunk sections
            chunk_table = doc.add_table(rows=1, cols=4)
            chunk_table.style = 'Light List Accent 1'
            
            hdr = chunk_table.rows[0].cells
            hdr[0].text = "Section"
            hdr[1].text = "Status"
            hdr[2].text = "Confidence"
            hdr[3].text = "Details / Reason"

            for sec in sections_in_chunk:
                name = sec.get("name", "Unknown")
                status_val = (sec.get("status", "") or "").lower()
                reason = sec.get("reason", "No details provided.")
                confidence = sec.get("confidence", "medium")
                
                row = chunk_table.add_row().cells
                row[0].text = name
                row[1].text = "✅ Pass" if status_val == "pass" else "❌ Fail"
                row[2].text = confidence.capitalize()
                row[3].text = reason
        else:
            doc.add_paragraph("⚠️ No sections identified in this chunk.")

        # Critical issues for this chunk
        chunk_critical = chunk.get("critical_issues", [])
        if chunk_critical:
            doc.add_paragraph("🚨 Critical Issues in This Chunk:", style="Heading 4")
            for issue in chunk_critical:
                p = doc.add_paragraph(issue, style="List Bullet")
                for run in p.runs:
                    run.font.color.rgb = RGBColor(239, 68, 68)

        doc.add_paragraph()  # Spacing

    # --- Footer ---
    doc.add_paragraph()
    footer = doc.add_paragraph("━" * 60)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    footer_text = doc.add_paragraph(
        "Powered by Azure GPT-4.1-mini + Enhanced Validation Engine | "
        "Generated by BRD Validation Assistant"
    )
    footer_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer_text.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(107, 114, 128)

    # --- Save to BytesIO ---
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
with tab8:
    st.markdown("""
    <style>
        .validator-header {
            text-align: center;
            padding: 0.8rem;
            background: linear-gradient(90deg, #10b981, #14b8a6);
            color: white;
            border-radius: 10px;
            font-size: 1.4rem;
            font-weight: 600;
            margin-bottom: 1.2rem;
        }
        .validator-box {
            background-color: #f9fafc;
            border: 1px solid #e5e7eb;
            border-radius: 12px;
            padding: 1.5rem;
            box-shadow: 0 4px 8px rgba(0,0,0,0.05);
        }
    </style>
    """, unsafe_allow_html=True)

    st.markdown('<div class="validator-header">🧾 BRD Validation Assistant</div>', unsafe_allow_html=True)

    st.markdown('<div class="validator-box">', unsafe_allow_html=True)
    uploaded_brd = st.file_uploader("📤 Upload your BRD (Word, PDF, or Text)", type=["docx", "pdf", "txt", "md"])
    validate_button = st.button("✅ Validate BRD", use_container_width=True)
    st.markdown('</div>', unsafe_allow_html=True)

    if validate_button and uploaded_brd:
        from brd_generator import validate_uploaded_brd
        import docx, tempfile, fitz, io, json

        # --- Extract text from uploaded file ---
        text = ""
        if uploaded_brd.name.endswith(".docx"):
            doc = docx.Document(uploaded_brd)
            text = "\n".join(p.text for p in doc.paragraphs)
        elif uploaded_brd.name.endswith(".pdf"):
            pdf = fitz.open(stream=uploaded_brd.read(), filetype="pdf")
            text = "\n".join(page.get_text() for page in pdf)
        else:
            text = uploaded_brd.read().decode("utf-8", errors="ignore")

        with st.spinner("🔍 Validating entire BRD (this may take a minute)..."):
            idx, texts, meta, _, _ = load_index(index_path)
            seed_results = query(idx, texts, meta,
                                "overview main modules data access business processes controllers",
                                top_k=15)
            result = validate_uploaded_brd(text, seed_results)

            # --- Display summary ---
            if result["status"].lower() == "valid":
                st.success("✅ BRD aligns with the application context.")
            else:
                st.error("❌ BRD has mismatches with the analyzed system.")

            st.markdown(f"**Summary:** {result['summary']}")
            st.divider()

            # --- Section-wise checks ---
            # st.subheader("📋 Section Results")
            # for key, val in result.get("checks", {}).items():
            #     icon = "✔️" if val.lower() == "pass" else "❌"
            #     st.write(f"{icon} **{key.replace('_',' ').title()}**")
            # --- Section-wise checks (dynamic) ---
            st.subheader("📋 Section Results")

            sections = result.get("sections", {})
            if not sections:
                st.info("No section-level results available.")
            else:
                for name, info in sections.items():
                    status = (info.get("status", "") or "").lower()
                    reason = info.get("reason", "") or ""
                    icon = "✔️" if status == "pass" else "❌"

                    if status == "pass":
                        st.write(f"{icon} **{name}** – Passed")
                    else:
                        st.write(f"{icon} **{name}** – Failed")
                        if reason:
                            st.caption(f"• {reason}")

            # --- Download detailed report ---
            st.divider()
            report_json = json.dumps(result, indent=2)
            st.download_button(
                label="📥 Download Detailed Validation Report (JSON)",
                data=report_json,
                file_name="brd_validation_report.json",
                mime="application/json",
                use_container_width=True
            )

            # --- Word report download ---
        from brd_generator import generate_brd_validation_docx
        word_data = generate_brd_validation_docx(result, seed_results)

        st.download_button(
            label="📘 Download BRD Validation Report (Word)",
            data=word_data,
            file_name="brd_validation_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )


from utils import chunk_document_safe
import json, re, time
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
from datetime import datetime

def validate_uploaded_brd(brd_text: str, seed_results: list) -> dict:
    """
    ✅ Enhanced BRD validation with dynamic section detection and detailed failure reasons.
    Validates uploaded BRD (supports 100+ pages) against repository insights.
    Uses shared _chat() for Azure GPT-4.1-mini.
    """

    # --- Context summary (limit safely) ---
    context_summary = "\n\n".join([r.get("text", "")[:2000] for r in seed_results[:10]])

    # --- Chunk entire BRD (~9k tokens each ≈ 20–25 pages) ---
    brd_chunks = chunk_document_safe(brd_text, max_tokens=9000)
    total_chunks = len(brd_chunks)
    print(f"🧾 Validating BRD ({total_chunks} chunks, full coverage)...")

    all_results = []

    for i, chunk in enumerate(brd_chunks, start=1):
        prompt = f"""
You are a senior BRD validator with expertise in software requirements analysis.

Compare this BRD content chunk ({i}/{total_chunks}) with the provided repository context.

**CRITICAL INSTRUCTIONS:**
1. IDENTIFY all sections present in this chunk (e.g., Executive Summary, Objectives, Scope, 
   Assumptions, Constraints, Risks, Functional Requirements, Non-Functional Requirements, 
   User Stories, Use Cases, APIs, Data Models, Security Requirements, etc.)
   
2. For EACH section found:
   - Determine if it ALIGNS with the application context
   - If it FAILS, provide SPECIFIC, ACTIONABLE reasons explaining:
     * What information is misaligned or missing
     * Which specific requirements conflict with the system
     * What technical gaps exist
     * Concrete examples of discrepancies

3. Respond ONLY in valid JSON format with this exact structure:

{{
  "status": "valid" or "invalid",
  "summary": "Brief 1-2 sentence evaluation of this chunk's alignment",
  "sections": [
    {{
      "name": "Exact section name from BRD",
      "status": "pass" or "fail",
      "reason": "Detailed, specific explanation (MANDATORY for fail status). Include concrete examples and technical details.",
      "confidence": "high" or "medium" or "low"
    }}
  ],
  "detected_sections_count": <number of sections identified>,
  "critical_issues": ["List of critical misalignments if any"]
}}

**VALIDATION RULES:**
- Only include sections that ACTUALLY appear in this chunk
- For PASS: Confirm alignment with specific details
- For FAIL: Provide detailed, actionable reasons with examples
- Use "confidence" to indicate validation certainty
- Flag "critical_issues" for major problems (security gaps, conflicting requirements, etc.)

--- REPOSITORY CONTEXT ---
{context_summary}

--- BRD CHUNK ({i}/{total_chunks}) ---
{chunk[:8500]}

Remember: Be thorough, specific, and provide actionable feedback for any failures.
"""

        # --- Retry logic using _chat() ---
        for attempt in range(3):
            try:
                response_text = _chat([
                    {"role": "system", "content": "You are an enterprise BRD validator. Respond only in valid JSON format."},
                    {"role": "user", "content": prompt}
                ], temperature=0.1)

                # Extract JSON from response
                json_match = re.search(r"\{.*\}", response_text, re.DOTALL)
                if not json_match:
                    raise ValueError("No JSON object found in model output")

                parsed = json.loads(json_match.group(0))

                # Normalize structure
                if "sections" not in parsed or not isinstance(parsed["sections"], list):
                    parsed["sections"] = []
                if "critical_issues" not in parsed:
                    parsed["critical_issues"] = []
                if "detected_sections_count" not in parsed:
                    parsed["detected_sections_count"] = len(parsed["sections"])

                parsed["chunk_id"] = i
                all_results.append(parsed)
                print(f"✓ Chunk {i}/{total_chunks} validated ({parsed['detected_sections_count']} sections found)")
                break
                
            except Exception as e:
                print(f"⚠️ Chunk {i}/{total_chunks} failed (attempt {attempt+1}): {e}")
                time.sleep(2 + attempt)
                if attempt == 2:
                    # Fallback entry when the model fails 3 times
                    all_results.append({
                        "chunk_id": i,
                        "status": "invalid",
                        "summary": f"Chunk {i} failed validation due to model or parsing errors after 3 attempts.",
                        "sections": [{
                            "name": "Validation Error",
                            "status": "fail",
                            "reason": f"Technical error during validation: {str(e)}. Unable to process this section of the BRD.",
                            "confidence": "low"
                        }],
                        "detected_sections_count": 0,
                        "critical_issues": ["Validation process failure"]
                    })

    if not all_results:
        return {
            "status": "invalid",
            "summary": "No valid responses from validation model.",
            "sections": {},
            "chunks": [],
            "critical_issues": ["Complete validation failure"]
        }

    # --- Aggregate Results with Enhanced Tracking ---
    section_stats = {}
    valid_count = 0
    invalid_count = 0
    all_critical_issues = []

    for r in all_results:
        # Track chunk-level status
        if r.get("status", "invalid").lower() == "valid":
            valid_count += 1
        else:
            invalid_count += 1

        # Collect critical issues
        critical = r.get("critical_issues", [])
        if critical:
            all_critical_issues.extend(critical)

        # Aggregate section-level results
        for sec in r.get("sections", []):
            name = (sec.get("name") or "").strip()
            if not name:
                continue
                
            status = (sec.get("status") or "fail").lower()
            reason = (sec.get("reason") or "").strip()
            confidence = sec.get("confidence", "medium")

            if name not in section_stats:
                section_stats[name] = {
                    "pass": 0,
                    "fail": 0,
                    "fail_reasons": [],
                    "pass_confirmations": [],
                    "confidence_scores": []
                }

            section_stats[name]["confidence_scores"].append(confidence)
            
            if status == "pass":
                section_stats[name]["pass"] += 1
                if reason:
                    section_stats[name]["pass_confirmations"].append(reason)
            else:
                section_stats[name]["fail"] += 1
                if reason:
                    section_stats[name]["fail_reasons"].append(reason)

    # --- Build Final Aggregated Sections with Detailed Reasons ---
    aggregated_sections = {}
    for name, stats in section_stats.items():
        is_pass = stats["pass"] >= stats["fail"]
        status = "pass" if is_pass else "fail"

        # Determine confidence
        confidence_map = {"high": 3, "medium": 2, "low": 1}
        avg_confidence_score = sum(confidence_map.get(c, 2) for c in stats["confidence_scores"]) / len(stats["confidence_scores"])
        if avg_confidence_score >= 2.5:
            confidence = "high"
        elif avg_confidence_score >= 1.5:
            confidence = "medium"
        else:
            confidence = "low"

        # Build detailed reason
        reason = ""
        if status == "fail":
            # Deduplicate and format failure reasons
            unique_reasons = list(dict.fromkeys(stats["fail_reasons"]))
            if unique_reasons:
                reason = " | ".join(unique_reasons[:5])  # Show up to 5 distinct reasons
            else:
                reason = "Multiple alignment issues detected across chunks. Review detailed chunk analysis for specifics."
        else:
            # For passing sections, optionally include confirmation
            if stats["pass_confirmations"]:
                reason = stats["pass_confirmations"][0]  # Use first confirmation as summary

        aggregated_sections[name] = {
            "status": status,
            "reason": reason,
            "confidence": confidence,
            "pass_count": stats["pass"],
            "fail_count": stats["fail"]
        }

    # --- Overall Status and Summary ---
    overall_status = "valid" if valid_count >= invalid_count else "invalid"
    total_sections = len(aggregated_sections)
    failed_sections = sum(1 for s in aggregated_sections.values() if s["status"] == "fail")
    
    summary = (
        f"Validated {total_chunks} chunks covering {total_sections} distinct sections. "
        f"{valid_count} chunks passed, {invalid_count} failed. "
        f"{failed_sections} sections require attention."
    )

    # Add critical issues to summary if any
    if all_critical_issues:
        unique_critical = list(dict.fromkeys(all_critical_issues))[:3]
        summary += f" Critical issues: {', '.join(unique_critical)}"

    aggregated = {
        "status": overall_status,
        "summary": summary,
        "sections": aggregated_sections,
        "chunks": all_results,
        "critical_issues": list(dict.fromkeys(all_critical_issues)),
        "statistics": {
            "total_chunks": total_chunks,
            "valid_chunks": valid_count,
            "invalid_chunks": invalid_count,
            "total_sections": total_sections,
            "passed_sections": total_sections - failed_sections,
            "failed_sections": failed_sections
        }
    }

    print(f"✅ BRD Validation Completed: {total_sections} sections across {total_chunks} chunks")
    return aggregated


def generate_brd_validation_docx(result: dict, seed_results: list) -> bytes:
    """
    📄 Generate an enhanced, professional Word report (.docx) for BRD validation.
    Includes dynamic checklist, detailed failure reasons, and visual formatting.
    Returns raw bytes ready for Streamlit download.
    """

    doc = Document()

    # --- Cover Page ---
    title = doc.add_heading("BRD Validation Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(16, 185, 129)  # Green color

    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", style="Body Text").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Validated using Azure GPT-4.1-mini against repository insights", style="Body Text").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # --- Executive Summary ---
    doc.add_heading("📊 Executive Summary", level=1)
    
    status = (result.get("status", "unknown") or "").upper()
    status_text = "✅ VALID" if status == "VALID" else "❌ INVALID"
    status_para = doc.add_paragraph()
    status_run = status_para.add_run(f"Overall Validation Status: {status_text}")
    status_run.bold = True
    status_run.font.size = Pt(12)
    if status == "VALID":
        status_run.font.color.rgb = RGBColor(16, 185, 129)
    else:
        status_run.font.color.rgb = RGBColor(239, 68, 68)

    doc.add_paragraph(result.get("summary", "No summary available."))

    # Statistics
    stats = result.get("statistics", {})
    if stats:
        doc.add_paragraph()
        doc.add_paragraph("📈 Validation Statistics:", style="Heading 3")
        doc.add_paragraph(f"• Total Chunks Analyzed: {stats.get('total_chunks', 0)}")
        doc.add_paragraph(f"• Valid Chunks: {stats.get('valid_chunks', 0)}")
        doc.add_paragraph(f"• Invalid Chunks: {stats.get('invalid_chunks', 0)}")
        doc.add_paragraph(f"• Total Sections Found: {stats.get('total_sections', 0)}")
        doc.add_paragraph(f"• Passed Sections: {stats.get('passed_sections', 0)}")
        doc.add_paragraph(f"• Failed Sections: {stats.get('failed_sections', 0)}")

    # Critical Issues
    critical_issues = result.get("critical_issues", [])
    if critical_issues:
        doc.add_paragraph()
        doc.add_heading("⚠️ Critical Issues", level=2)
        for issue in critical_issues[:10]:  # Limit to 10
            p = doc.add_paragraph(issue, style="List Bullet")
            for run in p.runs:
                run.font.color.rgb = RGBColor(239, 68, 68)

    doc.add_page_break()

    # --- Dynamic Section Checklist ---
    doc.add_heading("✅ BRD Section Validation Checklist", level=1)
    doc.add_paragraph("This checklist is dynamically generated based on sections detected in your BRD.")

    sections = result.get("sections", {})
    if sections:
        # Create table with headers
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Section Name"
        hdr_cells[1].text = "Status"
        hdr_cells[2].text = "Confidence"
        hdr_cells[3].text = "Occurrences"
        hdr_cells[4].text = "Validation Details / Reason"

        # Make headers bold
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # Add section data
        for name, info in sorted(sections.items()):
            row_cells = table.add_row().cells
            
            # Section name
            row_cells[0].text = name
            
            # Status with icon
            status_val = (info.get("status", "") or "").lower()
            status_text = "✅ Pass" if status_val == "pass" else "❌ Fail"
            row_cells[1].text = status_text
            
            # Confidence
            confidence = info.get("confidence", "medium")
            row_cells[2].text = confidence.capitalize()
            
            # Occurrences
            pass_count = info.get("pass_count", 0)
            fail_count = info.get("fail_count", 0)
            row_cells[3].text = f"✓{pass_count} / ✗{fail_count}"
            
            # Reason/Details
            reason = info.get("reason", "") or "No specific details provided."
            row_cells[4].text = reason
            
            # Color coding for failed sections
            if status_val == "fail":
                for paragraph in row_cells[1].paragraphs:
                    for run in paragraph.runs:
                        run.font.color.rgb = RGBColor(239, 68, 68)

        # Summary after table
        doc.add_paragraph()
        passed = sum(1 for s in sections.values() if s.get("status") == "pass")
        failed = sum(1 for s in sections.values() if s.get("status") == "fail")
        doc.add_paragraph(f"Summary: {passed} sections passed, {failed} sections failed out of {len(sections)} total sections detected.")
        
    else:
        doc.add_paragraph("⚠️ No sections were detected in the uploaded BRD. Please verify the document format.")

    doc.add_page_break()

    # --- Repository Context Reference ---
    doc.add_heading("📚 Repository Context Summary", level=1)
    doc.add_paragraph("The following repository insights were used as validation baseline:")
    
    for i, r in enumerate(seed_results[:8], start=1):
        snippet = (r.get("text", "") or "").strip().replace("\n", " ")
        if snippet:
            doc.add_paragraph(f"{i}. {snippet[:600]}{'...' if len(snippet) > 600 else ''}", style="List Number")

    doc.add_page_break()

    # --- Detailed Chunk Analysis ---
    doc.add_heading("🔍 Detailed Chunk-Level Analysis", level=1)
    doc.add_paragraph("Below is a detailed breakdown of validation results for each document chunk:")

    for chunk in result.get("chunks", []):
        chunk_id = chunk.get("chunk_id", "?")
        chunk_status = chunk.get("status", "unknown")
        
        # Chunk header
        heading = doc.add_heading(f"Chunk {chunk_id} - {chunk_status.upper()}", level=2)
        if chunk_status.lower() == "invalid":
            for run in heading.runs:
                run.font.color.rgb = RGBColor(239, 68, 68)
        
        doc.add_paragraph(chunk.get("summary", "No summary available."))
        
        sections_in_chunk = chunk.get("sections", [])
        detected_count = chunk.get("detected_sections_count", len(sections_in_chunk))
        doc.add_paragraph(f"Sections detected in this chunk: {detected_count}")

        if sections_in_chunk:
            # Create mini-table for chunk sections
            chunk_table = doc.add_table(rows=1, cols=4)
            chunk_table.style = 'Light List Accent 1'
            
            hdr = chunk_table.rows[0].cells
            hdr[0].text = "Section"
            hdr[1].text = "Status"
            hdr[2].text = "Confidence"
            hdr[3].text = "Details / Reason"

            for sec in sections_in_chunk:
                name = sec.get("name", "Unknown")
                status_val = (sec.get("status", "") or "").lower()
                reason = sec.get("reason", "No details provided.")
                confidence = sec.get("confidence", "medium")
                
                row = chunk_table.add_row().cells
                row[0].text = name
                row[1].text = "✅ Pass" if status_val == "pass" else "❌ Fail"
                row[2].text = confidence.capitalize()
                row[3].text = reason
        else:
            doc.add_paragraph("⚠️ No sections identified in this chunk.")

        # Critical issues for this chunk
        chunk_critical = chunk.get("critical_issues", [])
        if chunk_critical:
            doc.add_paragraph("🚨 Critical Issues in This Chunk:", style="Heading 4")
            for issue in chunk_critical:
                p = doc.add_paragraph(issue, style="List Bullet")
                for run in p.runs:
                    run.font.color.rgb = RGBColor(239, 68, 68)

        doc.add_paragraph()  # Spacing

    # --- Footer ---
    doc.add_paragraph()
    footer = doc.add_paragraph("━" * 60)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    footer_text = doc.add_paragraph(
        "Powered by Azure GPT-4.1-mini + Enhanced Validation Engine | "
        "Generated by BRD Validation Assistant"
    )
    footer_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer_text.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(107, 114, 128)

    # --- Save to BytesIO ---
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()





def add_documents_to_index(
    index_path: str,
    new_docs: List[Dict[str, Any]],
    use_checkpoint: bool = True
) -> None:
    """
    Add new documents to existing index with validation.
    Now supports Excel (.xls/.xlsx) supporting docs in the same way
    as build_faiss_index (per-sheet chunking + rich metadata).
    """
    print(f"\n➕ Adding {len(new_docs)} documents to existing index...")

    # Load existing index + metadata
    index, texts, metadata, _, _ = load_faiss_index(index_path)

    new_texts: List[str] = []
    new_metadata: List[Dict[str, Any]] = []

    def _sliding_window_text_simple(text: str, size: int = 1800, overlap_chars: int = 200) -> List[str]:
        step = max(1, size - overlap_chars)
        out = []
        for i in range(0, len(text), step):
            out.append(text[i:i + size])
            if i + size >= len(text):
                break
        return out

    print("📄 Normalizing new documents (including Excel if present)...")

    for doc in new_docs:
        # Figure out filename / extension
        file_name = doc.get("source") or doc.get("title") or doc.get("file_name") or "document"
        file_ext = (doc.get("ext") or "").lower()

        # Heuristic to detect Excel supporting docs (same as build_faiss_index)
        is_excel = (
            file_ext in (".xlsx", ".xls")
            or str(file_name).lower().endswith((".xlsx", ".xls"))
            or doc.get("file_bytes")
            or doc.get("content_bytes")
            or (doc.get("file_path") and str(doc.get("file_path")).lower().endswith((".xlsx", ".xls")))
        )

        # === Excel path ======================================================
        if is_excel:
            print(f"📊 Detected Excel document for append: {file_name}")

            # If upload already split it into per-sheet entries (type == 'excel_sheet'),
            # or we only have workbook bytes, _process_excel_doc_to_chunks handles both.
            excel_chunks = _process_excel_doc_to_chunks(doc)

            if not excel_chunks:
                # Fallback: if there's plain text, at least index that instead of skipping
                doc_text = doc.get("text", "")
                if doc_text and doc_text.strip():
                    # Prefer chunk_document_safe if available
                    try:
                        from utils import chunk_document_safe
                        doc_chunks = chunk_document_safe(doc_text, max_tokens=7500)
                    except Exception:
                        doc_chunks = _sliding_window_text_simple(doc_text)

                    for i, piece in enumerate(doc_chunks):
                        if not piece.strip():
                            continue
                        new_texts.append(piece)
                        new_metadata.append({
                            "type": doc.get("type", "document"),
                            "title": f"{file_name} (Part {i+1}/{len(doc_chunks)})",
                            "source": file_name,
                            "sheet": None,
                            "text": piece,
                            "chunk_id": i
                        })
                else:
                    print(f"⚠️ Skipping Excel doc with no readable content: {file_name}")
                continue  # done with this doc

            # Normal Excel handling (match build_faiss_index metadata)
            for c in excel_chunks:
                txt = c.get("text", "")
                if not txt or not txt.strip():
                    continue

                new_texts.append(txt)
                new_metadata.append({
                    "type": c.get("type", "support_doc_table"),
                    "title": c.get("title"),
                    "file_name": c.get("file_name"),
                    "source": file_name,
                    "sheet": c.get("sheet_name") or c.get("sheet"),
                    "sheet_name": c.get("sheet_name") or c.get("sheet"),
                    "sheet_rows": c.get("sheet_rows"),
                    "sheet_cols": c.get("sheet_cols"),
                    "headers": c.get("headers"),
                    "text": txt,
                    "chunk_id": c.get("chunk_id"),
                    # keep CSV bytes so that the exact sheet can be reconstructed later
                    "sheet_csv_bytes": c.get("sheet_csv_bytes")
                })

            continue  # move on to next doc

        # === Non-Excel documents ============================================
        doc_text = doc.get("text", "")
        if not doc_text or not doc_text.strip():
            print(f"⚠️  Skipping empty document: {file_name}")
            continue

        # Prefer the same safe chunking strategy as build_faiss_index
        try:
            from utils import chunk_document_safe
            doc_chunks = chunk_document_safe(doc_text, max_tokens=7500)
        except Exception:
            doc_chunks = _sliding_window_text_simple(doc_text)

        for i, piece in enumerate(doc_chunks):
            if not piece.strip():
                continue
            new_texts.append(piece)
            new_metadata.append({
                "type": doc.get("type", "document"),
                "title": f"{file_name} (Part {i+1}/{len(doc_chunks)})",
                "source": file_name,
                "sheet": None,
                "text": piece,
                "chunk_id": i
            })

    if not new_texts:
        print("⚠️  No valid text segments to add!")
        return

    print(f"✅ Prepared {len(new_texts)} new text segments for embedding")

    # Generate embeddings with checkpoint
    checkpoint_path = None
    if use_checkpoint:
        Path(CHECKPOINT_DIR).mkdir(parents=True, exist_ok=True)
        checkpoint_path = os.path.join(CHECKPOINT_DIR, "append_checkpoint.pkl")

    try:
        print(f"\n📊 Generating embeddings for {len(new_texts)} new segments...")
        new_embeddings = _embed_texts_with_retry(new_texts, checkpoint_path=checkpoint_path)

        # Clear checkpoint on success
        if checkpoint_path:
            _clear_checkpoint(checkpoint_path)

    except Exception as e:
        print(f"❌ Failed to generate embeddings for new documents: {e}")

        if checkpoint_path and Path(checkpoint_path).exists():
            print(f"\n💾 Checkpoint saved at: {checkpoint_path}")
            print(f"💡 Run the append operation again to resume from checkpoint.")

        raise

    if len(new_embeddings) == 0:
        print("⚠️  No embeddings generated for new documents!")
        return

    # Add to index using FAISS
    print(f"\n➕ Adding {len(new_embeddings)} vectors to existing FAISS index...")
    index.add(new_embeddings.astype("float32"))

    # Update metadata list in memory
    metadata.extend(new_metadata)

    # Save updated index
    faiss.write_index(index, index_path)

    meta_path = index_path.replace(".faiss", ".meta.pkl")
    with open(meta_path, "wb") as f:
        pickle.dump(metadata, f)

    print(f"✅ Updated index now has {index.ntotal} vectors")
    print(f"📝 Metadata entries now: {len(metadata)}")



with tab1:
    st.header("📁 Source Code Upload & Analysis")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.subheader("Upload Repository")
        up_zip = st.file_uploader(
            "Upload a .zip of your repository", 
            type=["zip"], 
            accept_multiple_files=False
        )
        
        if up_zip:
            try:
                z = zipfile.ZipFile(io.BytesIO(up_zip.read()))
                z.extractall(workdir)
                st.success(f"✅ ZIP extracted to: {workdir}")
                
                # Show extracted files
                extracted_files = []
                for root, dirs, files in os.walk(workdir):
                    for file in files:
                        rel_path = os.path.relpath(os.path.join(root, file), workdir)
                        if any(rel_path.endswith(ext) for ext in CODE_EXTS):      
                            extracted_files.append(rel_path)
                
                st.info(f"Found {len(extracted_files)} code files")
                with st.expander("View extracted files"):
                    for f in sorted(extracted_files)[:50]:  # Show first 50
                        st.text(f)
                    if len(extracted_files) > 50:
                        st.text(f"... and {len(extracted_files) - 50} more files")
                        
            except Exception as e:
                st.error(f"Error extracting ZIP: {str(e)}")
    
    with col2:
        st.subheader("Upload Individual Files")
        up_files = st.file_uploader(
            "...or upload individual code files", 
            type=[e.lstrip(".") for e in CODE_EXTS], 
            accept_multiple_files=True
        )
        
        if up_files:
            try:
                for f in up_files:
                    dest = os.path.join(workdir, f.name)
                    os.makedirs(os.path.dirname(dest), exist_ok=True)
                    with open(dest, "wb") as out:
                        out.write(f.read())
                st.success(f"✅ Saved {len(up_files)} files")
            except Exception as e:
                st.error(f"Error saving files: {str(e)}")
    
    # === NEW SECTION: Supporting Documents ===
    # === NEW SECTION: Supporting Documents (UPDATED to add Excel bytes + richer metadata) ===
    st.divider()
    st.subheader("📎 Upload Supporting Documents (optional)")

    doc_types = ["pdf", "docx", "xlsx", "csv", "txt", "md"]
    support_docs = st.file_uploader(
        "Upload additional documents such as BRDs, specs, or references",
        type=doc_types,
        accept_multiple_files=True
    )

    if support_docs:
        try:
            from docx import Document as DocxDocument
            import PyPDF2
            import pandas as pd
            from utils import chunk_document_safe

            for doc_file in support_docs:
                file_name = doc_file.name
                file_ext = os.path.splitext(file_name)[1].lower()

                # Initialize container
                doc_text = ""
                file_bytes = None

                # Read bytes for binary formats (xlsx, pdf, docx) to allow later parsing
                try:
                    file_bytes = doc_file.read()
                except Exception:
                    # If reading bytes fails, attempt to get text/stream directly
                    try:
                        doc_file.seek(0)
                        file_bytes = doc_file.getvalue()
                    except Exception:
                        file_bytes = None

                # Extract text based on file type (lightweight summary for immediate UI)
                if file_ext == '.txt' or file_ext == '.md':
                    try:
                        doc_text = file_bytes.decode('utf-8', errors='ignore') if file_bytes else ''
                    except Exception:
                        doc_text = ''

                elif file_ext == '.docx':
                    try:
                        docx = DocxDocument(io.BytesIO(file_bytes))
                        doc_text = '\n\n'.join([para.text for para in docx.paragraphs if para.text.strip()])
                    except Exception as e:
                        print(f"⚠️ Failed to parse DOCX {file_name}: {e}")
                        doc_text = ''

                elif file_ext == '.pdf':
                    try:
                        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
                        pages = [page.extract_text() or "" for page in pdf_reader.pages]
                        doc_text = '\n\n'.join(pages)
                    except Exception as e:
                        print(f"⚠️ Failed to parse PDF {file_name}: {e}")
                        doc_text = ''

                elif file_ext == '.csv':
                    try:
                        df = pd.read_csv(io.BytesIO(file_bytes))
                        # keep a concise summary for UI; full table bytes saved for indexing
                        doc_text = df.head(20).to_string()
                    except Exception as e:
                        print(f"⚠️ Failed to parse CSV {file_name}: {e}")
                        doc_text = ''

                # elif file_ext == '.xlsx' or file_ext == '.xls':
                #     # For Excel, save raw bytes so indexer can parse multi-sheet structure later.
                #     try:
                #         # Lightweight preview: list sheet names and first few rows of first sheet
                #         xls = pd.ExcelFile(io.BytesIO(file_bytes))
                #         sheet_names = xls.sheet_names
                #         preview = []
                #         if sheet_names:
                #             first = pd.read_excel(io.BytesIO(file_bytes), sheet_name=sheet_names[0])
                #             preview.append(f"Sheets: {sheet_names}")
                #             preview.append("First sheet preview:\n" + first.head(10).to_string())
                #             doc_text = '\n\n'.join(preview)
                #         else:
                #             doc_text = ''
                #     except Exception as e:
                #         print(f"⚠️ Failed to parse Excel preview {file_name}: {e}")
                #         doc_text = ''

                elif file_ext == '.xlsx' or file_ext == '.xls':
                    try:
                        # Parse all sheets and create one support_doc per sheet (preserve structure)
                        xls = pd.ExcelFile(io.BytesIO(file_bytes))
                        sheet_names = xls.sheet_names
                        workbook_title = file_name
                        st.info(f"📄 {file_name}: Excel uploaded ({len(sheet_names)} sheets)")

                        # Create a top-level workbook metadata entry
                        st.session_state.support_docs.append({
                            **metadata_base,
                            "type": "excel_workbook",
                            "sheets": sheet_names,
                            "file_bytes": file_bytes,  # keep raw bytes if needed later
                            "text": "",  # no flattened text at workbook level
                            "chunk_id": 0
                        })

                        for sheet in sheet_names:
                            try:
                                df = pd.read_excel(io.BytesIO(file_bytes), sheet_name=sheet)
                                if df is None:
                                    continue

                                rows, cols = getattr(df, "shape", (0, 0))
                                headers = list(df.columns.astype(str))

                                # create CSV preview text (for indexing/embedding)
                                csv_preview = df.head(200).to_csv(index=False)  # preview first 200 rows
                                # store full sheet as CSV bytes for later reconstruction (can be large)
                                full_csv_bytes = df.to_csv(index=False).encode("utf-8")

                                st.session_state.support_docs.append({
                                    **metadata_base,
                                    "type": "excel_sheet",
                                    "sheet_name": sheet,
                                    "sheet_rows": rows,
                                    "sheet_cols": cols,
                                    "headers": headers,
                                    "text": csv_preview,
                                    "sheet_csv_bytes": full_csv_bytes,  # full sheet CSV bytes
                                    "file_bytes": None,  # workbook bytes already stored above if needed
                                    "chunk_id": 0
                                })

                            except Exception as e:
                                print(f"⚠️ Failed to parse sheet {sheet} in {file_name}: {e}")
                                continue

                    except Exception as e:
                        print(f"⚠️ Failed to parse Excel preview {file_name}: {e}")
                        doc_text = ''

                else:
                    # Unknown extension — attempt generic text extraction
                    try:
                        doc_text = file_bytes.decode('utf-8', errors='ignore') if file_bytes else ''
                    except Exception:
                        doc_text = ''

                # Chunk large text (summary/user-visible) for immediate UI and quick-search
                if doc_text and doc_text.strip():
                    chunks = chunk_document_safe(doc_text, max_tokens=7500)
                else:
                    chunks = []

                # Initialize support_docs in session_state
                if 'support_docs' not in st.session_state:
                    st.session_state.support_docs = []

                # If we have raw bytes for Excel, include them so the indexer can fully parse sheets
                metadata_base = {
                    'title': file_name,
                    'source': file_name,
                    'ext': file_ext,
                    'type': 'document'
                }

                if file_bytes and file_ext in ('.xlsx', '.xls'):
                    # Store bytes along with a lightweight text preview (if any)
                    st.session_state.support_docs.append({
                        **metadata_base,
                        'file_bytes': file_bytes,
                        'text': doc_text or '',
                        'chunk_id': 0
                    })
                    st.info(f"📄 {file_name}: Excel uploaded ({len(xls.sheet_names) if 'xls' in locals() else 'unknown'} sheets)")

                    # Also store any preview chunks for immediate searching / UI
                    for i, chunk in enumerate(chunks):
                        st.session_state.support_docs.append({
                            **metadata_base,
                            'text': chunk,
                            'chunk_id': i
                        })

                else:
                    # Non-Excel docs: store chunked text representations (as before)
                    if chunks:
                        st.info(f"📄 {file_name}: {len(''.join(chunks))} chars → {len(chunks)} chunk(s)")
                        for i, chunk in enumerate(chunks):
                            st.session_state.support_docs.append({
                                **metadata_base,
                                'text': chunk,
                                'chunk_id': i
                            })
                    else:
                        # if no text extracted, still store bytes if available (e.g., images or unknown)
                        if file_bytes:
                            st.session_state.support_docs.append({
                                **metadata_base,
                                'file_bytes': file_bytes,
                                'text': doc_text or '',
                                'chunk_id': 0
                            })
                            st.warning(f"⚠️ {file_name}: No readable text extracted; saved raw bytes for later processing.")
                        else:
                            st.warning(f"⚠️ Could not extract text from {file_name}")

            st.success(f"✅ Processed {len(support_docs)} supporting documents")

        except Exception as e:
            st.error(f"Error processing documents: {str(e)}")
            st.exception(e)

    # --- rest of the UI (notes, parse buttons) remains unchanged ---

    st.markdown("### ✍️ Add Supporting Notes or Context")
    user_notes = st.text_area(
        "You can write additional context, notes, or descriptions here to support repository analysis.",
        placeholder="e.g., This repository contains the backend services for our NLP system...",
        height=200
    )

    if user_notes.strip():
        from utils import chunk_document_safe
        chunks = chunk_document_safe(user_notes, max_tokens=7500)

        
        if "support_docs" not in st.session_state:
            st.session_state.support_docs = []
        
        for i, chunk in enumerate(chunks):
            st.session_state.support_docs.append({
                "text": chunk,
                "title": f"Manual Note (Part {i+1}/{len(chunks)})",
                "source": "User Input",
                "type": "manual_note",
                "chunk_id": i
            })
        
        st.success(f"✅ Added {len(chunks)} chunk(s) from manual input")

    st.divider()
    
    # Parse repository
    col1, col2 = st.columns([3, 1])
    with col1:
        st.subheader("🔍 Repository Analysis")
    with col2:
        parse_btn = st.button("🚀 Parse Repository", type="primary")
    print(f"Append mode: {st.session_state.append_mode}")
    if parse_btn and not st.session_state.append_mode:
        # Check if there's already parsed data
        if "parsed" in st.session_state and st.session_state["parsed"]:
            st.warning("⚠️ Note: Parsing will replace current analysis. To ADD to existing index, use 'Append Mode' in sidebar after parsing.")
        
        with st.spinner("Analyzing repository..."):
            try:
                parsed = parse_repository_enhanced(workdir)
                                                #    , max_chunk = 1800, overlap = 200)
                
                # Add supporting documents to parsed object if they exist
                if "support_docs" in st.session_state and st.session_state.support_docs:
                    parsed["supporting_docs"] = st.session_state.support_docs
                    st.info(f"📎 Including {len(st.session_state.support_docs)} supporting document chunks")
                else:
                    parsed["supporting_docs"] = []
                
                st.session_state["parsed"] = parsed
                
                # Quick summary
                col1, col2, col3, col4 = st.columns(4)
                with col1:
                    st.metric("Code Files", parsed["metrics"]["total"]["total_files"])
                with col2:
                    st.metric("Lines of Code", f"{parsed['metrics']['total']['total_loc']:,}")
                with col3:
                    st.metric("Components", len(parsed["nodes"]))
                with col4:
                    st.metric("Processes", len(parsed["business_processes"]))
                
                st.success("✅ Repository parsed successfully!")
                st.balloons()
                # st.snow()
                
                # Store in graph database (optional)
                if gremlin_enabled and GRAPH_STORE_AVAILABLE:
                    try:
                        gs = GraphStore()
                        gs.upsert_vertices(parsed["nodes"])
                        gs.upsert_edges(parsed["edges"])
                        st.info("📊 Graph data stored in Cosmos DB")
                    except Exception as e:
                        st.warning(f"Graph storage failed: {str(e)}")
                elif gremlin_enabled and not GRAPH_STORE_AVAILABLE:
                    st.warning("⚠️ Graph storage configured but graph_store.py not found")
                
            except Exception as e:
                st.error(f"Parsing failed: {str(e)}")
    elif parse_btn and st.session_state.append_mode:
        print("Append mode parsing...")
        parsed = parse_repository_enhanced(workdir)
        def concat_results(res1: dict, res2: dict) -> dict:
            def merge_field(v1, v2):
                # If both are lists — concatenate
                if isinstance(v1, list) and isinstance(v2, list):
                    return v1 + v2
                # If both are dicts — merge keys (res2 overwrites res1 on conflict)
                elif isinstance(v1, dict) and isinstance(v2, dict):
                    return {**v1, **v2}
                # If one is None — return the other
                elif v1 is None:
                    return v2
                elif v2 is None:
                    return v1
                # Otherwise — prefer first or handle conflict
                else:
                    return v1

            return {
                "nodes": res1.get("nodes", []) + res2.get("nodes", []),
                "edges": res1.get("edges", []) + res2.get("edges", []),
                "chunks": res1.get("chunks", []) + res2.get("chunks", []),
                "metrics": {
                    "total": merge_field(res1.get("metrics", {}).get("total"), res2.get("metrics", {}).get("total")),
                    "by_file": merge_field(res1.get("metrics", {}).get("by_file"), res2.get("metrics", {}).get("by_file"))
                },
                "business_processes": merge_field(res1.get("business_processes"), res2.get("business_processes")),
                "power_platform_mapping": merge_field(res1.get("power_platform_mapping"), res2.get("power_platform_mapping")),
                "comprehensive_schema": merge_field(res1.get("comprehensive_schema"), res2.get("comprehensive_schema")),
                "supporting_docs": res1.get("supporting_docs", []) + res2.get("supporting_docs", [])
            }

        final = concat_results(st.session_state["parsed"] , parsed)
        print(f"Final parsed nodes: {len(final['nodes'])}, edges: {len(final['edges'])}, chunks: {len(final['chunks'])}")
        st.session_state["parsed"] = final
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("Code Files", final["metrics"]["total"]["total_files"])
        with col2:
            st.metric("Lines of Code", f"{final['metrics']['total']['total_loc']:,}")
        with col3:
            st.metric("Components", len(final["nodes"]))
        with col4:
            st.metric("Processes", len(final["business_processes"]))
        
        st.success("✅ Repository parsed successfully!")
        st.balloons()






def _process_excel_doc_to_chunks(doc: Dict[str, Any], max_chunk_chars: int = 1800, overlap: int = 200) -> List[Dict[str, Any]]:
    """
    Given a supporting_doc entry that represents an Excel workbook or sheet,
    return a list of normalized document chunks (each with 'text' and metadata fields).
    Accepts doc entries that are:
      - type == 'excel_sheet' (already parsed per-sheet during upload), OR
      - have 'file_bytes' for a workbook and need to be parsed here.
    """
    chunks = []
    file_name = doc.get("source") or doc.get("title") or doc.get("file_name") or "excel_document"
    file_ext = (doc.get("ext") or "").lower()

    # If doc already represents a single sheet (upload step created per-sheet entries), use it directly
    if doc.get("type") == "excel_sheet":
        sheet_name = doc.get("sheet_name")
        sheet_rows = doc.get("sheet_rows")
        sheet_cols = doc.get("sheet_cols")
        headers = doc.get("headers", [])
        # prefer existing preview text for embeddings/indexing
        sheet_text = doc.get("text", "")
        if not sheet_text.strip() and doc.get("sheet_csv_bytes"):
            try:
                sheet_text = doc["sheet_csv_bytes"].decode("utf-8", errors="ignore")
            except Exception:
                sheet_text = ""

        if sheet_text:
            # Use chunk_document_safe if available
            try:
                from utils import chunk_document_safe
                sub_chunks = chunk_document_safe(sheet_text, max_tokens=7500)
            except Exception:
                # fallback sliding window
                def sliding_window_text(text, size=max_chunk_chars, step=None):
                    if step is None:
                        step = max(1, size - overlap)
                    out = []
                    for i in range(0, max(1, len(text)), step):
                        out.append(text[i:i+size])
                        if i + size >= len(text):
                            break
                    return out
                sub_chunks = sliding_window_text(sheet_text, max_chunk_chars, max_chunk_chars - overlap)

            for i, piece in enumerate(sub_chunks):
                chunks.append({
                    "text": piece,
                    "type": "support_doc_table",
                    "title": f"{file_name} - {sheet_name} (Part {i+1}/{len(sub_chunks)})",
                    "file_name": file_name,
                    "sheet": sheet_name,
                    "sheet_name": sheet_name,
                    "sheet_rows": sheet_rows,
                    "sheet_cols": sheet_cols,
                    "headers": headers,
                    "chunk_id": i,
                    # keep original full bytes in metadata if present (useful later)
                    "sheet_csv_bytes": doc.get("sheet_csv_bytes")
                })
        return chunks

    # Otherwise, fall back to parsing bytes (workbook) and produce per-sheet chunks
    excel_bytes = None
    if doc.get("file_bytes"):
        excel_bytes = doc["file_bytes"]
    elif doc.get("content_bytes"):
        excel_bytes = doc["content_bytes"]
    elif doc.get("file_path") and os.path.exists(doc["file_path"]):
        with open(doc["file_path"], "rb") as f:
            excel_bytes = f.read()

    if not excel_bytes:
        return chunks

    # Parse workbook into sheets
    try:
        sheets = pd.read_excel(io.BytesIO(excel_bytes), sheet_name=None)
    except Exception:
        try:
            sheets = pd.read_excel(io.BytesIO(excel_bytes), sheet_name=None, engine="openpyxl")
        except Exception as e:
            print(f"⚠️ Failed to parse excel for {file_name}: {e}")
            return chunks

    for sheet_name, df in sheets.items():
        if df is None or df.empty:
            continue
        try:
            sheet_text = df.to_csv(index=False)
        except Exception:
            sheet_text = df.to_string(index=False)

        rows, cols = getattr(df, "shape", (None, None))
        headers = list(df.columns.astype(str))

        # chunk sheet_text (prefer chunk_document_safe)
        try:
            from utils import chunk_document_safe
            sub_chunks = chunk_document_safe(sheet_text, max_tokens=7500)
        except Exception:
            # fallback sliding window
            def sliding_window_text(text, size=max_chunk_chars, step=None):
                if step is None:
                    step = max(1, size - overlap)
                out = []
                for i in range(0, max(1, len(text)), step):
                    out.append(text[i:i+size])
                    if i + size >= len(text):
                        break
                return out
            sub_chunks = sliding_window_text(sheet_text, max_chunk_chars, max_chunk_chars - overlap)

        full_csv_bytes = sheet_text.encode("utf-8")
        for i, piece in enumerate(sub_chunks):
            chunks.append({
                "text": piece,
                "type": "support_doc_table",
                "title": f"{file_name} - {sheet_name} (Part {i+1}/{len(sub_chunks)})",
                "file_name": file_name,
                "sheet": sheet_name,
                "sheet_name": sheet_name,
                "sheet_rows": rows,
                "sheet_cols": cols,
                "headers": headers,
                "chunk_id": i,
                "sheet_csv_bytes": full_csv_bytes
            })

    return chunks

def build_faiss_index(
    chunks: List[Dict[str, Any]],
    index_path: str,
    use_checkpoint: bool = True,
    supporting_docs: Optional[List[Dict[str, Any]]] = None
) -> Tuple[faiss.Index, List[Dict[str, Any]]]:
    """
    Build FAISS index from code/document chunks and optional supporting documents.
    This version supports Excel (.xls/.xlsx) supporting docs (multi-sheet),
    safe chunking, and richer metadata for sheet/table entries.
    """
    print(f"\n{'='*60}")
    print("🏗️  BUILDING FAISS INDEX (WITH SUPPORT DOCS)" if supporting_docs else "🏗️  BUILDING FAISS INDEX")
    print(f"{'='*60}")

    start_time = time.time()

    texts = []
    metadata = []

    print(f"📦 Processing {len(chunks)} code chunks...")
    for chunk in chunks:
        chunk_text = chunk.get("text", "")
        if not chunk_text or not chunk_text.strip():
            continue

        texts.append(chunk_text)
        metadata.append({
            "text": chunk_text,
            "type": chunk.get("type", "code_chunk"),
            "path": chunk.get("path", ""),
            "kind": chunk.get("kind", "code_chunk"),
            "hash": chunk.get("hash", ""),
            "metrics": chunk.get("metrics", {})
        })

    # Process supporting docs
    if supporting_docs:
        print(f"📄 Adding {len(supporting_docs)} supporting documents...")
        for doc in supporting_docs:
            # Determine if doc is excel by extension or presence of bytes/file_path
            file_name = doc.get("source") or doc.get("title") or doc.get("file_name") or "support_doc"
            file_ext = (doc.get("ext") or "").lower()
            # Heuristic: check ext or file_name suffix or presence of bytes
            is_excel = file_ext in (".xlsx", ".xls") or str(file_name).lower().endswith((".xlsx", ".xls")) \
                       or doc.get("file_bytes") or doc.get("content_bytes") or (doc.get("file_path") and str(doc.get("file_path")).lower().endswith((".xlsx", ".xls")))

            # if is_excel:
            #     excel_chunks = _process_excel_doc_to_chunks(doc)
            #     if not excel_chunks:
            #         # If parser couldn't extract sheets, fallback to any provided text
            #         if doc.get("text") and doc.get("text").strip():
            #             texts.append(doc.get("text"))
            #             metadata.append({
            #                 "type": doc.get("type", "document"),
            #                 "title": file_name,
            #                 "source": file_name,
            #                 "sheet": None,
            #                 "text": doc.get("text"),
            #                 "chunk_id": doc.get("chunk_id", 0)
            #             })
            #         else:
            #             print(f"⚠️ Skipping excel doc with no readable content: {file_name}")
            #         continue

            #     for c in excel_chunks:
            #         txt = c.get("text", "")
            #         if not txt or not txt.strip():
            #             continue
            #         texts.append(txt)
            #         metadata.append({
            #             "type": c.get("type", "support_doc_table"),
            #             "title": c.get("title"),
            #             "file_name": c.get("file_name"),
            #             "source": file_name,
            #             "sheet": c.get("sheet"),
            #             "sheet_rows": c.get("sheet_rows"),
            #             "sheet_cols": c.get("sheet_cols"),
            #             "text": txt,
            #             "chunk_id": c.get("chunk_id")
            #         })
            
            if is_excel:
                # If doc already has per-sheet info (created at upload), _process_excel_doc_to_chunks handles it
                excel_chunks = _process_excel_doc_to_chunks(doc)
                if not excel_chunks:
                    # fallback: if doc.text exists, index that
                    if doc.get("text") and doc.get("text").strip():
                        texts.append(doc.get("text"))
                        metadata.append({
                            "type": doc.get("type", "document"),
                            "title": file_name,
                            "source": file_name,
                            "sheet": None,
                            "text": doc.get("text"),
                            "chunk_id": doc.get("chunk_id", 0)
                        })
                    else:
                        print(f"⚠️ Skipping excel doc with no readable content: {file_name}")
                    continue

                for c in excel_chunks:
                    txt = c.get("text", "")
                    if not txt or not txt.strip():
                        continue
                    texts.append(txt)
                    # preserve structured metadata so you can reconstruct the exact sheet later
                    metadata.append({
                        "type": c.get("type", "support_doc_table"),
                        "title": c.get("title"),
                        "file_name": c.get("file_name"),
                        "source": file_name,
                        "sheet": c.get("sheet_name") or c.get("sheet"),
                        "sheet_name": c.get("sheet_name") or c.get("sheet"),
                        "sheet_rows": c.get("sheet_rows"),
                        "sheet_cols": c.get("sheet_cols"),
                        "headers": c.get("headers"),
                        "text": txt,
                        "chunk_id": c.get("chunk_id"),
                        # persist sheet bytes for exact reconstruction if present
                        "sheet_csv_bytes": c.get("sheet_csv_bytes")
                    })

            
            
            else:
                # Not an excel file — use plain text already present or split larger text into pieces
                doc_text = doc.get("text", "")
                if doc_text and doc_text.strip():
                    # prefer chunk_document_safe if available
                    try:
                        from utils import chunk_document_safe
                        doc_chunks = chunk_document_safe(doc_text, max_tokens=7500)
                    except Exception:
                        # fallback to simple sliding window by characters
                        def sliding_window_text_simple(text, size=1800, overlap_chars=200):
                            step = max(1, size - overlap_chars)
                            out = []
                            for i in range(0, len(text), step):
                                out.append(text[i:i+size])
                                if i + size >= len(text):
                                    break
                            return out
                        doc_chunks = sliding_window_text_simple(doc_text)

                    for i, piece in enumerate(doc_chunks):
                        if not piece.strip():
                            continue
                        texts.append(piece)
                        metadata.append({
                            "type": doc.get("type", "document"),
                            "title": f"{file_name} (Part {i+1}/{len(doc_chunks)})",
                            "source": file_name,
                            "sheet": None,
                            "text": piece,
                            "chunk_id": i
                        })
                else:
                    print(f"⚠️  Skipping empty document: {file_name}")

    if not texts:
        raise ValueError("No valid chunks or documents to index!")

    print(f"✅ {len(texts)} total text segments prepared for embedding")

    # Generate embeddings (with checkpoint support)
    checkpoint_path = None
    if use_checkpoint:
        Path(CHECKPOINT_DIR).mkdir(parents=True, exist_ok=True)
        checkpoint_path = os.path.join(CHECKPOINT_DIR, "build_checkpoint.pkl")

    try:
        print(f"\n📊 Generating embeddings for {len(texts)} texts...")
        embeddings = _embed_texts_with_retry(texts, checkpoint_path=checkpoint_path)

        if len(embeddings) == 0:
            raise ValueError("No embeddings generated!")

        if checkpoint_path:
            _clear_checkpoint(checkpoint_path)

    except Exception as e:
        print(f"\n❌ Failed to generate embeddings: {e}")
        if checkpoint_path and Path(checkpoint_path).exists():
            print(f"\n💾 Checkpoint saved at: {checkpoint_path}")
            print("💡 Run the index build again to resume from checkpoint.")
        raise

    # Create FAISS index
    print(f"\n🔍 Creating FAISS index...")
    dim = embeddings.shape[1]
    index = faiss.IndexFlatL2(dim)
    index.add(embeddings.astype('float32'))

    # Save index + metadata
    print(f"💾 Saving index to {index_path}...")
    Path(index_path).parent.mkdir(parents=True, exist_ok=True)
    faiss.write_index(index, index_path)

    meta_path = index_path.replace(".faiss", ".meta.pkl")
    print(f"💾 Saving metadata to {meta_path}...")
    with open(meta_path, "wb") as f:
        pickle.dump(metadata, f)

    total_time = time.time() - start_time
    print(f"\n{'='*60}")
    print(f"✅ INDEX BUILD COMPLETE")
    print(f"{'='*60}")
    print(f"   📊 Total vectors: {index.ntotal}")
    print(f"   📐 Vector dimension: {dim}")
    print(f"   📝 Metadata entries: {len(metadata)}")
    print(f"   🕒 Total time: {total_time:.2f} seconds")
    print(f"{'='*60}\n")

    print(f"🔎 Verifying FAISS index and metadata alignment...")

    if index.ntotal != len(metadata):
        print(f"⚠️  Mismatch detected! FAISS vectors: {index.ntotal}, Metadata entries: {len(metadata)}")
    else:
        print(f"✅ FAISS vectors ({index.ntotal}) perfectly match metadata entries.")

    # Confirm metadata file exists and is readable
    if not Path(meta_path).exists():
        print(f"❌ Metadata file not found at {meta_path}")
    else:
        try:
            with open(meta_path, "rb") as f:
                loaded_meta = pickle.load(f)
            if len(loaded_meta) == len(metadata):
                print(f"✅ Metadata file successfully verified ({len(loaded_meta)} entries)")
            else:
                print(f"⚠️ Metadata file entry count mismatch ({len(loaded_meta)} vs {len(metadata)})")
        except Exception as e:
            print(f"❌ Failed to read metadata file: {e}")

    return index, metadata

